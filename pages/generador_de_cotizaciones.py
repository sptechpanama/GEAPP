from __future__ import annotations

import base64
import html
import json
import uuid
import os
import re
import math
import time
import logging
import zipfile
import unicodedata
from datetime import date, datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

import pandas as pd
import requests
import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import load_workbook
from openpyxl.drawing.image import Image
from openpyxl.styles import Alignment, Border, Side
from gspread.exceptions import WorksheetNotFound
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload
from sheets import get_client, read_worksheet, write_worksheet
from entities import client_selector, _load_clients, WS_CLIENTES
from ui.theme import apply_global_theme

st.set_page_config(page_title="Generador de cotizaciones", page_icon="🧾", layout="wide")
apply_global_theme()

# ---- Guard simple ----
if st.session_state.get("authentication_status") is not True:
    st.switch_page("Inicio.py")


# ---- Helpers ----
def _load_logo_b64(*paths: str) -> str:
    for path in paths:
        if not path:
            continue
        try:
            with open(path, "rb") as fh:
                return base64.b64encode(fh.read()).decode()
        except Exception:
            continue
    return ""


def _format_money(value: float) -> str:
    return f"${value:,.2f}"


def _slugify_text(text: str, *, max_words: int = 8, max_len: int = 56) -> str:
    raw = str(text or "").strip().lower()
    if not raw:
        return ""
    normalized = unicodedata.normalize("NFKD", raw)
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
    slug = re.sub(r"[^a-z0-9]+", "-", ascii_text).strip("-")
    if not slug:
        return ""
    words = [w for w in slug.split("-") if w]
    if max_words > 0:
        words = words[:max_words]
    short_slug = "-".join(words).strip("-")
    if max_len > 0:
        short_slug = short_slug[:max_len].strip("-")
    return short_slug


def _quote_excel_filename(numero_cot: str, descripcion_corta: str) -> str:
    base = str(numero_cot or "").strip() or "COTIZACION"
    suffix = _slugify_text(descripcion_corta, max_words=8, max_len=56)
    if suffix:
        return f"{base}-{suffix}.xlsx"
    return f"{base}.xlsx"


def _openai_api_key() -> str:
    candidates: list[str | None] = []
    try:
        app_cfg = st.secrets.get("app", {})
        candidates.append(app_cfg.get("OPENAI_API_KEY"))
    except Exception:
        pass
    try:
        candidates.append(st.secrets.get("OPENAI_API_KEY"))
    except Exception:
        pass
    candidates.append(os.environ.get("OPENAI_API_KEY"))
    for raw in candidates:
        if raw and str(raw).strip():
            return str(raw).strip()
    return ""


def _openai_model_name() -> str:
    candidates: list[str | None] = []
    try:
        app_cfg = st.secrets.get("app", {})
        candidates.append(app_cfg.get("OPENAI_MODEL"))
        candidates.append(app_cfg.get("OPENAI_CHAT_MODEL"))
    except Exception:
        pass
    try:
        candidates.append(st.secrets.get("OPENAI_MODEL"))
    except Exception:
        pass
    candidates.append(os.environ.get("OPENAI_MODEL"))
    for raw in candidates:
        if raw and str(raw).strip():
            return str(raw).strip()
    return "gpt-4o-mini"


def _call_openai_chat(
    *,
    api_key: str,
    model: str,
    messages: list[dict[str, str]],
    temperature: float,
    max_tokens: int,
) -> str:
    response = requests.post(
        "https://api.openai.com/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json={
            "model": model,
            "messages": messages,
            "temperature": float(temperature),
            "max_tokens": int(max_tokens),
        },
        timeout=45,
    )
    response.raise_for_status()
    payload = response.json()
    choices = payload.get("choices") or []
    if not choices:
        raise RuntimeError("OpenAI no devolvió contenido.")
    content = choices[0].get("message", {}).get("content", "")
    if isinstance(content, list):
        out = []
        for item in content:
            if isinstance(item, dict):
                out.append(str(item.get("text", "")))
            else:
                out.append(str(item))
        return "".join(out).strip()
    return str(content).strip()


def _enforce_short_description_words(value: str, *, fallback_items: Optional[list[str]] = None) -> str:
    text = re.sub(r"[\r\n\t]+", " ", str(value or "").strip())
    text = re.sub(r"[`\"']", "", text)
    text = re.sub(r"\s+", " ", text).strip(" .,:;|-")
    words = [w for w in text.split(" ") if w]
    banned_words = {
        "cotizacion",
        "cotizaciones",
        "cliente",
        "empresa",
        "engineering",
        "medical",
        "rs",
        "rir",
        "panama",
        "compra",
        "estandar",
        "pc",
    }
    words = [w for w in words if w.lower() not in banned_words]
    if len(words) > 8:
        words = words[:8]

    if len(words) < 3 and fallback_items:
        extra_tokens: list[str] = []
        for item in fallback_items:
            item_text = re.sub(r"[\r\n\t]+", " ", str(item or "").strip())
            item_text = re.sub(r"[`\"']", "", item_text)
            item_text = re.sub(r"\s+", " ", item_text).strip(" .,:;|-")
            for token in item_text.split(" "):
                token = token.strip()
                if not token:
                    continue
                if token.lower() in banned_words:
                    continue
                extra_tokens.append(token)
                if len(extra_tokens) >= 8:
                    break
            if len(extra_tokens) >= 8:
                break
        for token in extra_tokens:
            if len(words) >= 8:
                break
            if token not in words:
                words.append(token)

    if not words:
        return "insumos generales de laboratorio"
    return " ".join(words)


def _fallback_quote_description(
    *,
    tipo_cotizacion: str,
    cliente: str,
    detalles: str,
    items: list[str],
) -> str:
    cliente_short = " ".join(str(cliente or "").split()[:3]).strip()
    if items:
        item_short = " ".join(str(items[0]).split()[:8]).strip()
        candidate = item_short
    elif detalles:
        detalle_short = " ".join(str(detalles).split()[:8]).strip()
        candidate = detalle_short
    else:
        tipo_short = str(tipo_cotizacion or "").strip().lower()
        candidate = tipo_short if tipo_short else "insumos generales de laboratorio"
    if cliente_short and not items and not detalles:
        candidate = f"{candidate} {cliente_short}".strip()
    return _enforce_short_description_words(candidate, fallback_items=items)


def _generate_quote_short_description(
    *,
    tipo_cotizacion: str,
    empresa: str,
    cliente: str,
    detalles: str,
    items: list[str],
) -> str:
    fallback = _fallback_quote_description(
        tipo_cotizacion=tipo_cotizacion,
        cliente=cliente,
        detalles=detalles,
        items=items,
    )
    api_key = _openai_api_key()
    if not api_key:
        return fallback
    try:
        model = _openai_model_name()
        items_text = ", ".join([str(x).strip() for x in items if str(x).strip()][:5])
        prompt = (
            "Genera una descripcion corta del OBJETO DEL ACTO/COMPRA.\n"
            "Reglas obligatorias:\n"
            "- Solo 3 a 8 palabras.\n"
            "- Español.\n"
            "- Sin comillas ni explicaciones.\n"
            "- Sin punto final.\n\n"
            "- Debe describir producto/insumo/equipo del acto.\n"
            "- No mencionar cliente, empresa ni tipo de cotización.\n"
            "- No usar palabras: cotizacion, cliente, empresa, rs, rir, panama, compra, estandar.\n\n"
            f"Items: {items_text or '-'}\n"
            f"Detalles: {(detalles or '-')[:220]}"
        )
        raw = _call_openai_chat(
            api_key=api_key,
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": "Responde solo con una frase corta sobre el objeto del acto.",
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
            max_tokens=40,
        )
        return _enforce_short_description_words(raw, fallback_items=items)
    except Exception:
        return fallback


def _extract_item_names_from_row(row: dict, limit: int = 5) -> list[str]:
    names: list[str] = []
    try:
        items = json.loads(row.get("items_json") or "[]")
        if isinstance(items, list):
            for item in items:
                if isinstance(item, dict):
                    val = str(item.get("producto_servicio") or item.get("descripcion") or "").strip()
                    if val:
                        names.append(val)
                if len(names) >= limit:
                    break
    except Exception:
        pass
    if not names:
        resumen = str(row.get("items_resumen") or "").strip()
        if resumen:
            parts = [p.strip() for p in resumen.split("|") if p.strip()]
            for part in parts[:limit]:
                names.append(part[:80])
    return names[:limit]


def _generate_description_for_row(row: dict) -> str:
    return _generate_quote_short_description(
        tipo_cotizacion=str(row.get("tipo_cotizacion") or ""),
        empresa=str(row.get("empresa") or ""),
        cliente=str(row.get("cliente_nombre") or ""),
        detalles=str(row.get("detalles_extra") or ""),
        items=_extract_item_names_from_row(row),
    )


SHEET_NAME_COT = "cotizaciones"
COT_COLUMNS = [
    "id",
    "numero_cotizacion",
    "prefijo",
    "secuencia",
    "empresa",
    "tipo_cotizacion",
    "descripcion_corta",
    "cliente_nombre",
    "cliente_direccion",
    "cliente_ruc",
    "cliente_dv",
    "fecha_cotizacion",
    "created_at",
    "updated_at",
    "moneda",
    "subtotal",
    "impuesto_pct",
    "impuesto_monto",
    "total",
    "items_json",
    "items_resumen",
    "detalles_extra",
    "presupuesto_items_json",
    "presupuesto_subtotal",
    "presupuesto_factor_ganancia",
    "presupuesto_precio_cotizar",
    "presupuesto_ganancia",
    "presupuesto_financiamiento_tipo",
    "presupuesto_financiamiento_interes_pct",
    "presupuesto_costo_financiamiento",
    "presupuesto_ganancia_neta",
    "presupuesto_t_inversion_presentacion",
    "presupuesto_inversion_etapa_1",
    "presupuesto_t_inicio_ejecucion_presentacion",
    "presupuesto_inversion_etapa_intermedia",
    "presupuesto_t_presentacion_cobro",
    "presupuesto_inversion_etapa_2",
    "presupuesto_t_recuperacion",
    "condiciones_json",
    "vigencia",
    "forma_pago",
    "entrega",
    "lugar_entrega",
    "estado",
    "notas",
    "drive_file_id",
    "drive_file_name",
    "drive_file_url",
    "drive_folder",
    "presupuesto_drive_file_id",
    "presupuesto_drive_file_name",
    "presupuesto_drive_file_url",
]
COT_PREFIX = {
    "RS Engineering": "RS",
    "RIR Medical": "RIR",
}
LP_DOC_TAB_NAME = "LP_ Doc_Generator"
LP_DOC_TIPO = "LP Doc Generator"
LP_DRIVE_FILENAME_BY_EMPRESA = {
    "RS Engineering": "LP_Doc_Generator_RS.xlsx",
    "RIR Medical": "LP_Doc_Generator_RIR.xlsx",
}
LP_BUDGET_FILENAME_BY_EMPRESA = {
    "RS Engineering": "LP_Doc_Generator_RS_Presupuesto.html",
    "RIR Medical": "LP_Doc_Generator_RIR_Presupuesto.html",
}
DEFAULT_COT_DRIVE_FOLDER_ID = "0AOB-QlptrUHYUk9PVA"
CLIENT_COLUMNS = ["RowID", "ClienteID", "ClienteNombre", "Empresa"]
CLIENT_EMPRESA_MAP = {
    "RS Engineering": "RS-SP",
    "RIR Medical": "RIR",
}
CLIENT_EMPRESA_OPTIONS = ["RS-SP", "RIR"]

PC_MANUAL_SHEET_ID_DEFAULT = "1-2sgJPhSPzP65HLeGSvxDBtfNczhiDiZhdEbyy6lia0"
PC_MANUAL_WORKSHEET = "pc_manual"
PC_CONFIG_WORKSHEET = "pc_config"
PC_CONFIG_HEADERS = ["name", "python", "script", "days", "times", "enabled"]
MANUAL_HEADERS = [
    "id",
    "job",
    "requested_by",
    "requested_at",
    "status",
    "notes",
    "payload",
    "result_file_id",
    "result_file_url",
    "result_file_name",
    "result_error",
]

ORQUESTADOR_JOB_NAME = "cotizacion_panama"
ORQUESTADOR_JOB_LABEL = "Cotizacion Panama Compra"
ORQUESTADOR_JOB_PY = r"C:\Users\rodri\scrapers_repo\.venv\Scripts\python.exe"
ORQUESTADOR_JOB_SCRIPT = r"C:\Users\rodri\selenium_cotizacion\cotizacion_worker.py"

GEAPP_ROOT = Path(__file__).resolve().parents[1]
SELENIUM_COTIZACION_DIR = Path(r"C:\Users\rodri\selenium_cotizacion")
REPO_COTIZACION_BASE_DIR = GEAPP_ROOT / "assets" / "cotizacion_base"


def _resolve_base_asset(file_name: str) -> Path:
    repo_candidate = REPO_COTIZACION_BASE_DIR / file_name
    if repo_candidate.exists():
        return repo_candidate
    return SELENIUM_COTIZACION_DIR / file_name


TEMPLATE_RS_STANDARD = _resolve_base_asset("plantilla_cotizacion.xlsx")
TEMPLATE_RIR_STANDARD = _resolve_base_asset("plantilla_cotizacion_rir.xlsx")
HEADER_RS_STANDARD = _resolve_base_asset("encabezado.png")
HEADER_RIR_STANDARD = _resolve_base_asset("encabezado_rir.png")
SIGNATURE_STANDARD = _resolve_base_asset("firma.png")
DOC_GEN_LOCAL_DIR = Path(r"C:\Users\rodri\doc_gen")
DOC_GEN_REPO_DIR = GEAPP_ROOT / "assets" / "doc_gen_base"
LP_DOC_FOLDER_NAME = "LP_Doc_Generator"
LP_TEMPLATE_FALLBACKS = {
    # Esta plantilla viene dañada desde origen (CRC en image1.tiff).
    # Fallback a la variante SF para no romper la generación.
    "template_no_incapacidad_para_contratar.docx": "template_no_incapacidad_para_contratar_sf.docx",
}
MESES_ES = {
    1: "enero",
    2: "febrero",
    3: "marzo",
    4: "abril",
    5: "mayo",
    6: "junio",
    7: "julio",
    8: "agosto",
    9: "septiembre",
    10: "octubre",
    11: "noviembre",
    12: "diciembre",
}

LP_DOC_SPECS = {
    "RS Engineering": [
        ("template_medidas_de_retorsion.docx", "medidas_de_retorsion.docx"),
        ("template_medidas_de_retorsion_sf.docx", "medidas_de_retorsion_sf.docx"),
        ("template_no_incapacidad_para_contratar.docx", "no_incapacidad_para_contratar.docx"),
        ("template_no_incapacidad_para_contratar_sf.docx", "no_incapacidad_para_contratar_sf.docx"),
        ("template_pacto_de_integridad.docx", "pacto_de_integridad.docx"),
        ("template_pacto_de_integridad_sf.docx", "pacto_de_integridad_sf.docx"),
        ("template_desglose_de_precios.docx", "desglose_de_precios.docx"),
        ("template_nota_adicional.docx", "nota_adicional.docx"),
        ("template_carta_de_adhesion.docx", "carta_de_adhesion.docx"),
        ("template_carta_de_adhesion_sf.docx", "carta_de_adhesion_sf.docx"),
    ],
    "RIR Medical": [
        ("template_medidas_de_retorsion_rir.docx", "medidas_de_retorsion.docx"),
        ("template_medidas_de_retorsion_sf_rir.docx", "medidas_de_retorsion_sf.docx"),
        ("template_no_incapacidad_para_contratar.docx", "no_incapacidad_para_contratar.docx"),
        ("template_no_incapacidad_para_contratar_sf.docx", "no_incapacidad_para_contratar_sf.docx"),
        ("template_pacto_de_integridad_rir.docx", "pacto_de_integridad.docx"),
        ("template_pacto_de_integridad_sf_rir.docx", "pacto_de_integridad_sf.docx"),
        ("template_desglose_de_precios_rir.docx", "desglose_de_precios.docx"),
        ("template_nota_adicional_rir.docx", "nota_adicional.docx"),
        ("template_carta_de_adhesion_rir.docx", "carta_de_adhesion.docx"),
        ("template_carta_de_adhesion_sf_rir.docx", "carta_de_adhesion_sf.docx"),
    ],
}


def _ensure_cotizaciones_sheet(client, sheet_id: str) -> None:
    sh = client.open_by_key(sheet_id)
    try:
        sh.worksheet(SHEET_NAME_COT)
        return
    except WorksheetNotFound:
        ws = sh.add_worksheet(title=SHEET_NAME_COT, rows=1000, cols=len(COT_COLUMNS))
        ws.update("A1", [COT_COLUMNS])


def _ensure_clientes_sheet(client, sheet_id: str) -> None:
    sh = client.open_by_key(sheet_id)
    try:
        sh.worksheet(WS_CLIENTES)
        return
    except WorksheetNotFound:
        ws = sh.add_worksheet(title=WS_CLIENTES, rows=1000, cols=len(CLIENT_COLUMNS))
        ws.update("A1", [CLIENT_COLUMNS])


def _pc_manual_sheet_id() -> str:
    try:
        app_cfg = st.secrets.get("app", {})
    except Exception:
        app_cfg = {}
    if isinstance(app_cfg, dict):
        return app_cfg.get("PC_MANUAL_SHEET_ID") or app_cfg.get("SHEET_ID") or PC_MANUAL_SHEET_ID_DEFAULT
    return PC_MANUAL_SHEET_ID_DEFAULT


def _pc_config_sheet_id() -> str:
    try:
        app_cfg = st.secrets.get("app", {})
    except Exception:
        app_cfg = {}
    if isinstance(app_cfg, dict):
        return (
            app_cfg.get("PC_CONFIG_SHEET_ID")
            or app_cfg.get("PC_MANUAL_SHEET_ID")
            or app_cfg.get("SHEET_ID")
            or PC_MANUAL_SHEET_ID_DEFAULT
        )
    return PC_MANUAL_SHEET_ID_DEFAULT


def _current_user() -> str:
    for key in ("username", "user", "email", "correo", "name", "nombre"):
        value = st.session_state.get(key)
        if value:
            return str(value)
    return "desconocido"


def _ensure_headers(client, sheet_id: str, worksheet: str, headers: list[str]) -> None:
    sh = client.open_by_key(sheet_id)
    try:
        ws = sh.worksheet(worksheet)
    except WorksheetNotFound:
        ws = sh.add_worksheet(title=worksheet, rows=200, cols=max(len(headers), 6))
        ws.update("A1", [headers])
        return

    existing = [cell.strip() for cell in (ws.row_values(1) or [])]
    if existing[: len(headers)] != headers:
        ws.update(f"A1:{chr(64 + len(headers))}1", [headers])


def _ensure_pc_config_job(client) -> None:
    sheet_id = _pc_config_sheet_id()
    _ensure_headers(client, sheet_id, PC_CONFIG_WORKSHEET, PC_CONFIG_HEADERS)
    sh = client.open_by_key(sheet_id)
    ws = sh.worksheet(PC_CONFIG_WORKSHEET)
    rows = ws.get_all_records() or []
    headers = [h.strip() for h in ws.row_values(1)]
    header_map = {h.lower(): idx + 1 for idx, h in enumerate(headers)}
    for idx, row in enumerate(rows, start=2):
        if str(row.get("name", "")).strip().lower() == ORQUESTADOR_JOB_NAME:
            updates = {
                "python": ORQUESTADOR_JOB_PY,
                "script": ORQUESTADOR_JOB_SCRIPT,
                "days": "",
                "times": "",
                "enabled": "si",
            }
            for key, value in updates.items():
                col = header_map.get(key)
                if not col:
                    continue
                current = str(row.get(key, "")).strip()
                if current != value:
                    ws.update_cell(idx, col, value)
            return

    row_data = {
        "name": ORQUESTADOR_JOB_NAME,
        "python": ORQUESTADOR_JOB_PY,
        "script": ORQUESTADOR_JOB_SCRIPT,
        "days": "",
        "times": "",
        "enabled": "si",
    }
    row_values = [row_data.get(col, "") for col in PC_CONFIG_HEADERS]
    ws.append_row(row_values, value_input_option="USER_ENTERED")


def _append_manual_request(client, payload: dict) -> str:
    sheet_id = _pc_manual_sheet_id()
    _ensure_headers(client, sheet_id, PC_MANUAL_WORKSHEET, MANUAL_HEADERS)
    sh = client.open_by_key(sheet_id)
    ws = sh.worksheet(PC_MANUAL_WORKSHEET)

    request_id = uuid.uuid4().hex
    row_data = {
        "id": request_id,
        "job": ORQUESTADOR_JOB_NAME,
        "requested_by": _current_user(),
        "requested_at": datetime.now(timezone.utc).astimezone().strftime("%Y-%m-%d %H:%M:%S"),
        "status": "pending",
        "notes": "",
        "payload": json.dumps(payload, ensure_ascii=False),
        "result_file_id": "",
        "result_file_url": "",
        "result_file_name": "",
        "result_error": "",
    }
    row_values = [row_data.get(header, "") for header in MANUAL_HEADERS]
    ws.append_row(row_values, value_input_option="USER_ENTERED")
    return request_id


def _fetch_manual_request(client, request_id: str) -> dict | None:
    sheet_id = _pc_manual_sheet_id()
    sh = client.open_by_key(sheet_id)
    ws = sh.worksheet(PC_MANUAL_WORKSHEET)
    values = ws.get_all_values()
    if not values:
        return None
    headers = [cell.strip() for cell in values[0]]
    for row in values[1:]:
        row_map = {headers[idx]: row[idx] if idx < len(row) else "" for idx in range(len(headers))}
        if str(row_map.get("id", "")).strip() == request_id:
            return row_map
    return None


def _extract_excel_items(excel_bytes: bytes) -> tuple[pd.DataFrame, str, bool, dict[str, Any]]:
    wb = load_workbook(BytesIO(excel_bytes))
    ws = wb["cotizacion"]

    def _to_float_safe(value) -> float:
        if value in (None, ""):
            return 0.0
        if isinstance(value, (int, float)):
            return float(value)
        text = str(value).strip()
        if not text:
            return 0.0
        text = (
            text.replace("B/.", "")
            .replace("B/", "")
            .replace("$", "")
            .replace(",", "")
            .strip()
        )
        text = re.sub(r"[^0-9.\-]", "", text)
        if text in {"", "-", ".", "-."}:
            return 0.0
        try:
            return float(text)
        except (TypeError, ValueError):
            return 0.0

    items = []
    row = 23
    while True:
        desc = ws[f"C{row}"].value
        unidad = ws[f"D{row}"].value
        cantidad = ws[f"E{row}"].value
        precio_unit = ws[f"F{row}"].value
        precio_total = ws[f"G{row}"].value

        marker = str(precio_unit or "").strip().lower()
        is_summary_row = (
            not any([desc, unidad, cantidad])
            and any(token in marker for token in ("subtotal", "impuesto", "total"))
        )
        if is_summary_row:
            break

        if not any([desc, unidad, cantidad, precio_unit, precio_total]):
            break

        cantidad_num = _to_float_safe(cantidad)
        precio_unit_num = _to_float_safe(precio_unit)
        precio_total_num = _to_float_safe(precio_total)
        if not str(desc or "").strip() and cantidad_num <= 0 and precio_unit_num <= 0 and precio_total_num <= 0:
            row += 1
            continue

        items.append(
            {
                "descripcion": desc or "",
                "unidad": unidad or "",
                "cantidad": cantidad_num,
                "precio_unitario": precio_unit_num,
                "precio_total": precio_total_num,
            }
        )
        row += 1
    titulo = str(ws["C19"].value or "").strip()
    itbms_row = 23 + len(items) + 1
    itbms_val = ws[f"G{itbms_row}"].value
    aplica_itbms = _to_float_safe(itbms_val) > 0
    entidad = str(ws["B13"].value or "").strip()
    fecha = str(ws["G13"].value or "").strip()
    numero_acto = str(ws["E18"].value or "").strip()
    ruc_dv = str(ws["B14"].value or "").strip()

    def _norm_label(value: str) -> str:
        raw = str(value or "").strip().lower()
        normalized = unicodedata.normalize("NFKD", raw)
        return normalized.encode("ascii", "ignore").decode("ascii")

    condiciones_map: dict[str, str] = {}
    for r in range(26, 140):
        text = str(ws[f"B{r}"].value or "").strip()
        if not text or ":" not in text:
            continue
        key, value = text.split(":", 1)
        key_norm = _norm_label(key)
        val = str(value or "").strip()
        if "forma de pago" in key_norm:
            condiciones_map["forma_pago"] = val
        elif "lugar de entrega" in key_norm:
            condiciones_map["lugar_entrega"] = val
        elif "tiempo de entrega" in key_norm:
            condiciones_map["tiempo_entrega"] = val
        elif "validez" in key_norm:
            condiciones_map["vigencia"] = val

    ruc = ""
    dv = ""
    m = re.search(r"RUC:\s*([^\s]+)\s+DV:\s*([^\s]+)", ruc_dv, flags=re.IGNORECASE)
    if m:
        ruc = str(m.group(1) or "").strip()
        dv = str(m.group(2) or "").strip()

    meta = {
        "entidad": entidad,
        "fecha": fecha,
        "numero_acto": numero_acto,
        "ruc": ruc,
        "dv": dv,
        "ruc_dv": ruc_dv,
        "forma_pago": condiciones_map.get("forma_pago", ""),
        "lugar_entrega": condiciones_map.get("lugar_entrega", ""),
        "tiempo_entrega": condiciones_map.get("tiempo_entrega", ""),
        "vigencia": condiciones_map.get("vigencia", ""),
    }
    wb.close()
    return pd.DataFrame(items), titulo, aplica_itbms, meta


def _apply_excel_edits(
    excel_bytes: bytes,
    items_df: pd.DataFrame,
    titulo: str,
    aplica_itbms: bool,
) -> bytes:
    wb = load_workbook(BytesIO(excel_bytes))
    ws = wb["cotizacion"]

    def _count_excel_items() -> int:
        row = 23
        count = 0
        while True:
            cells = [
                ws[f"C{row}"].value,
                ws[f"D{row}"].value,
                ws[f"E{row}"].value,
                ws[f"F{row}"].value,
                ws[f"G{row}"].value,
            ]
            if not any(cells):
                break
            count += 1
            row += 1
        return count
    items = items_df.copy()
    items["cantidad"] = pd.to_numeric(items.get("cantidad"), errors="coerce").fillna(0.0)
    items["precio_unitario"] = pd.to_numeric(items.get("precio_unitario"), errors="coerce").fillna(0.0)
    items["precio_total"] = items["cantidad"] * items["precio_unitario"]

    start_row = 23
    current_count = _count_excel_items()
    new_count = len(items)
    if new_count > current_count:
        ws.insert_rows(start_row + current_count, new_count - current_count)
    elif new_count < current_count and new_count > 0:
        ws.delete_rows(start_row + new_count, current_count - new_count)
    for idx, row in items.iterrows():
        target_row = start_row + idx
        ws[f"B{target_row}"] = idx + 1
        ws[f"C{target_row}"] = row.get("descripcion", "")
        ws[f"D{target_row}"] = row.get("unidad", "")
        ws[f"E{target_row}"] = float(row.get("cantidad", 0.0))
        ws[f"F{target_row}"] = float(row.get("precio_unitario", 0.0))
        ws[f"G{target_row}"] = float(row.get("precio_total", 0.0))

    if titulo:
        ws["C19"] = titulo
        ws["B21"] = titulo

    subtotal = float(items["precio_total"].sum())
    itbms = round(subtotal * 0.07, 2) if aplica_itbms else 0.0
    total = round(subtotal + itbms, 2)
    subtotal_row = start_row + len(items)
    itbms_row = subtotal_row + 1
    total_row = subtotal_row + 2
    ws[f"G{subtotal_row}"] = subtotal
    ws[f"G{itbms_row}"] = itbms
    ws[f"G{total_row}"] = total

    output = BytesIO()
    wb.save(output)
    wb.close()
    return output.getvalue()

@st.cache_data(show_spinner=False)
def _load_cotizaciones_cached(sheet_id: str, cache_token: str) -> pd.DataFrame:
    client, _ = get_client()
    _ensure_cotizaciones_sheet(client, sheet_id)
    df = read_worksheet(client, sheet_id, SHEET_NAME_COT)
    return df


def _normalize_cotizaciones_df(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    for col in COT_COLUMNS:
        if col not in out.columns:
            out[col] = ""
    out = out[COT_COLUMNS]
    for col in (
        "subtotal",
        "impuesto_pct",
        "impuesto_monto",
        "total",
        "presupuesto_subtotal",
        "presupuesto_factor_ganancia",
        "presupuesto_precio_cotizar",
        "presupuesto_ganancia",
        "presupuesto_financiamiento_interes_pct",
        "presupuesto_costo_financiamiento",
        "presupuesto_ganancia_neta",
        "presupuesto_t_inversion_presentacion",
        "presupuesto_inversion_etapa_1",
        "presupuesto_t_inicio_ejecucion_presentacion",
        "presupuesto_inversion_etapa_intermedia",
        "presupuesto_t_presentacion_cobro",
        "presupuesto_inversion_etapa_2",
        "presupuesto_t_recuperacion",
    ):
        out[col] = pd.to_numeric(out[col], errors="coerce")
    return out


def _normalize_clientes_df(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    for col in CLIENT_COLUMNS:
        if col not in out.columns:
            out[col] = ""
    out = out[CLIENT_COLUMNS]
    return out


def _next_sequence(df: pd.DataFrame, prefijo: str) -> int:
    if df.empty:
        return 1
    seq = pd.to_numeric(df.loc[df["prefijo"] == prefijo, "secuencia"], errors="coerce")
    if seq.dropna().empty:
        return 1
    return int(seq.max()) + 1


def _build_numero_cot(prefijo: str, secuencia: int) -> str:
    return f"COT-{prefijo}-{secuencia:04d}"


def _company_full_from_short(short_name: str) -> str:
    if str(short_name or "").strip().upper() == "RIR":
        return "RIR Medical"
    return "RS Engineering"


def _pc_prefijo_from_short(short_name: str) -> str:
    code = "RIR" if str(short_name or "").strip().upper() == "RIR" else "RS"
    return f"{code}-PC"


def _parse_manual_payload(raw_payload: str) -> dict[str, Any]:
    try:
        payload = json.loads(raw_payload or "{}")
        if isinstance(payload, dict):
            return payload
    except Exception:
        pass
    return {}


def _parse_panama_fecha(value: Any) -> date:
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    text = str(value or "").strip()
    if not text:
        return date.today()
    text = text[:10]
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y/%m/%d"):
        try:
            return datetime.strptime(text, fmt).date()
        except ValueError:
            continue
    return date.today()


def _build_standard_items_from_panama(raw_items_df: pd.DataFrame) -> pd.DataFrame:
    src = raw_items_df.copy()
    if "producto_servicio" not in src.columns:
        src["producto_servicio"] = src.get("descripcion", "")
    if "unidad" not in src.columns:
        src["unidad"] = src.get("unidad", "UND")
    if "cantidad" not in src.columns:
        src["cantidad"] = 0.0
    if "precio_unitario" not in src.columns:
        src["precio_unitario"] = 0.0
    if "precio_total" not in src.columns:
        src["precio_total"] = 0.0

    out = src[["producto_servicio", "unidad", "cantidad", "precio_unitario", "precio_total"]].copy()
    out["producto_servicio"] = out["producto_servicio"].fillna("").astype(str).str.strip()
    out["unidad"] = out["unidad"].fillna("UND").astype(str).str.strip().replace("", "UND")
    out["cantidad"] = pd.to_numeric(out["cantidad"], errors="coerce").fillna(0.0)
    out["precio_unitario"] = pd.to_numeric(out["precio_unitario"], errors="coerce").fillna(0.0)
    out["precio_total"] = pd.to_numeric(out["precio_total"], errors="coerce").fillna(0.0)

    # Fallback: si viene solo total por línea (sin cantidad/precio unitario), lo convertimos a item válido.
    missing_price = (out["precio_unitario"] <= 0) & (out["precio_total"] > 0)
    out.loc[missing_price & (out["cantidad"] <= 0), "cantidad"] = 1.0
    out.loc[missing_price, "precio_unitario"] = (
        out.loc[missing_price, "precio_total"] / out.loc[missing_price, "cantidad"].replace(0, 1.0)
    )
    out.loc[(out["producto_servicio"].str.len() == 0) & (out["precio_total"] > 0), "producto_servicio"] = (
        "Ítem Panamá Compra"
    )

    out = out[["producto_servicio", "unidad", "cantidad", "precio_unitario"]]
    out = out[
        (out["producto_servicio"].str.len() > 0)
        | (out["cantidad"] > 0)
        | (out["precio_unitario"] > 0)
    ].reset_index(drop=True)
    return out


def _create_cliente_in_sheet(
    client,
    sheet_id: str,
    nombre: str,
    empresa_codigo: str,
) -> tuple[str, bool]:
    _ensure_clientes_sheet(client, sheet_id)
    dfc = _normalize_clientes_df(read_worksheet(client, sheet_id, WS_CLIENTES))
    nombre_clean = nombre.strip()
    empresa_clean = empresa_codigo.strip()
    if not nombre_clean:
        raise ValueError("Debes indicar el nombre del cliente.")
    dup_mask = (
        dfc["ClienteNombre"].astype(str).str.strip().str.lower() == nombre_clean.lower()
    ) & (dfc["Empresa"].astype(str).str.strip().str.upper() == empresa_clean.upper())
    if not dfc.empty and dup_mask.any():
        row = dfc[dup_mask].iloc[0]
        return str(row.get("ClienteID", "")).strip(), False
    new_id = f"C-{uuid.uuid4().hex[:8].upper()}"
    new_row = {
        "RowID": uuid.uuid4().hex,
        "ClienteID": new_id,
        "ClienteNombre": nombre_clean,
        "Empresa": empresa_clean,
    }
    dfc = pd.concat([dfc, pd.DataFrame([new_row])], ignore_index=True)
    write_worksheet(client, sheet_id, WS_CLIENTES, dfc)
    _load_clients.clear()
    return new_id, True


def _get_drive_client(creds):
    return build("drive", "v3", credentials=creds)


def _find_or_create_folder(
    drive,
    name: str,
    parent_id: Optional[str] = None,
    drive_id: Optional[str] = None,
) -> str:
    query = ["mimeType='application/vnd.google-apps.folder'", "trashed=false", f"name='{name}'"]
    if parent_id:
        query.append(f"'{parent_id}' in parents")
    list_kwargs = {
        "q": " and ".join(query),
        "fields": "files(id,name)",
        "supportsAllDrives": True,
        "includeItemsFromAllDrives": True,
    }
    if drive_id:
        list_kwargs["corpora"] = "drive"
        list_kwargs["driveId"] = drive_id
    resp = drive.files().list(**list_kwargs).execute()
    files = resp.get("files", [])
    if files:
        return files[0]["id"]
    metadata = {"name": name, "mimeType": "application/vnd.google-apps.folder"}
    if parent_id:
        metadata["parents"] = [parent_id]
    created = drive.files().create(body=metadata, fields="id", supportsAllDrives=True).execute()
    return created["id"]


def _get_drive_folders(drive) -> tuple[str, Dict[str, str]]:
    base_id = st.secrets.get("app", {}).get("DRIVE_COTIZACIONES_FOLDER_ID") or DEFAULT_COT_DRIVE_FOLDER_ID
    drive_id = base_id
    subfolders = {
        "RS Engineering": _find_or_create_folder(drive, "RS", base_id, drive_id=drive_id),
        "RIR Medical": _find_or_create_folder(drive, "RIR", base_id, drive_id=drive_id),
    }
    return base_id, subfolders


def _find_file_in_folder(
    drive,
    *,
    folder_id: str,
    filename: str,
    drive_id: Optional[str] = None,
) -> Optional[str]:
    name = str(filename or "").replace("'", "\\'")
    if not folder_id or not name:
        return None
    query = (
        "trashed=false and "
        f"'{folder_id}' in parents and "
        f"name='{name}'"
    )
    list_kwargs = {
        "q": query,
        "fields": "files(id,name,modifiedTime)",
        "orderBy": "modifiedTime desc",
        "pageSize": 1,
        "supportsAllDrives": True,
        "includeItemsFromAllDrives": True,
    }
    if drive_id:
        list_kwargs["corpora"] = "drive"
        list_kwargs["driveId"] = drive_id
    resp = drive.files().list(**list_kwargs).execute()
    files = resp.get("files") or []
    if not files:
        return None
    return files[0].get("id")


def _resolve_doc_gen_template(file_name: str) -> Path:
    def _zip_health_error(path: Path) -> str:
        try:
            with zipfile.ZipFile(path, "r") as zf:
                bad_member = zf.testzip()
            if bad_member:
                return f"CRC inválido en {bad_member}"
            return ""
        except Exception as exc:  # pylint: disable=broad-except
            return str(exc)

    candidates = [
        DOC_GEN_REPO_DIR / file_name,
        DOC_GEN_LOCAL_DIR / file_name,
    ]
    first_error = ""
    for candidate in candidates:
        if not candidate.exists():
            continue
        health_error = _zip_health_error(candidate)
        if not health_error:
            return candidate
        if not first_error:
            first_error = f"{candidate}: {health_error}"
        fallback_name = LP_TEMPLATE_FALLBACKS.get(file_name)
        if fallback_name:
            fallback_candidate = candidate.with_name(fallback_name)
            if fallback_candidate.exists():
                fallback_error = _zip_health_error(fallback_candidate)
                if not fallback_error:
                    logging.warning(
                        "Plantilla %s dañada (%s). Usando fallback %s.",
                        candidate.name,
                        health_error,
                        fallback_candidate.name,
                    )
                    return fallback_candidate
    raise FileNotFoundError(
        f"No se encontró la plantilla de Doc_Generator: {file_name}. "
        f"Busca en {DOC_GEN_REPO_DIR} o {DOC_GEN_LOCAL_DIR}. "
        f"{('Detalle: ' + first_error) if first_error else ''}"
    )


def _replace_text_tokens(text: str, replacements: dict[str, str]) -> str:
    out = str(text or "")
    for key, value in replacements.items():
        out = out.replace(key, str(value or ""))
    return out


def _style_paragraph(paragraph, *, in_table: bool) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1 if in_table else 2)
    pf.line_spacing = 1.0
    text = str(paragraph.text or "").strip()
    if text:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if in_table else WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9.5 if in_table else 10.5)


def _apply_doc_professional_style(doc: Document) -> None:
    for section in doc.sections:
        # Forzar tamaño legal 8.5" x 14" y márgenes compactos.
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(14)
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.35)
        section.left_margin = Inches(0.45)
        section.right_margin = Inches(0.45)
        section.header_distance = Inches(0.2)
        section.footer_distance = Inches(0.2)

    for paragraph in doc.paragraphs:
        _style_paragraph(paragraph, in_table=False)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _style_paragraph(paragraph, in_table=True)


def _fill_docx_placeholders(template_path: Path, replacements: dict[str, str]) -> bytes:
    doc = Document(str(template_path))

    for paragraph in doc.paragraphs:
        new_text = _replace_text_tokens(paragraph.text, replacements)
        if new_text != paragraph.text:
            paragraph.text = new_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_text = _replace_text_tokens(paragraph.text, replacements)
                    if new_text != paragraph.text:
                        paragraph.text = new_text

    _apply_doc_professional_style(doc)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _fecha_spanish_tokens(fecha_base: date) -> dict[str, str]:
    mes = MESES_ES.get(int(fecha_base.month), str(fecha_base.month))
    return {
        "[fecha]": f"{fecha_base.day:02d} de {mes} de {fecha_base.year}",
        "[dia]": f"{fecha_base.day:02d}",
        "[mes]": mes,
        "[año]": str(fecha_base.year),
    }


def _extract_numero_acto_from_link(enlace: str) -> str:
    text = str(enlace or "").strip()
    if not text:
        return ""
    patterns = [
        r"/(?:solicitud-de-cotizacion|pliego-de-cargos)/([^/?#]+)/",
        r"/(?:solicitud-de-cotizacion|pliego-de-cargos)/([^/?#]+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return str(match.group(1) or "").strip()
    return ""


def _build_lp_doc_replacements(
    *,
    fecha_base: date,
    representante_legal_pacto: str,
    representante_legal_documentos: str,
    cedula: str,
    entidad: str,
    titulo: str,
    numero_acto: str,
    lugar_entrega: str,
    tiempo_entrega: str,
) -> dict[str, str]:
    replacements = _fecha_spanish_tokens(fecha_base)
    replacements.update(
        {
            "[Representante_legal_de_la_Entidad_Licitante]": str(
                representante_legal_documentos or ""
            ).strip(),
            "[entidad]": str(entidad or "").strip(),
            "[titulo]": str(titulo or "").strip(),
            "[numero_de_acto]": str(numero_acto or "").strip(),
            "[cedula]": str(cedula or "").strip(),
            "[lugar]": str(lugar_entrega or "").strip(),
            "[entrega]": str(tiempo_entrega or "").strip(),
        }
    )
    # En pacto de integridad Doc_Gen usa el representante legal "pacto".
    if str(representante_legal_pacto or "").strip():
        replacements["[Representante_legal_de_la_Entidad_Licitante]"] = str(
            representante_legal_pacto
        ).strip()
    return replacements


def _build_lp_documents(
    *,
    empresa_full: str,
    fecha_base: date,
    representante_legal_pacto: str,
    representante_legal_documentos: str,
    cedula: str,
    entidad: str,
    titulo: str,
    numero_acto: str,
    lugar_entrega: str,
    tiempo_entrega: str,
) -> list[dict[str, Any]]:
    specs = LP_DOC_SPECS.get(empresa_full) or LP_DOC_SPECS["RS Engineering"]
    docs_out: list[dict[str, Any]] = []
    for template_name, output_name in specs:
        replacements = _build_lp_doc_replacements(
            fecha_base=fecha_base,
            representante_legal_pacto=representante_legal_pacto,
            representante_legal_documentos=representante_legal_documentos,
            cedula=cedula,
            entidad=entidad,
            titulo=titulo,
            numero_acto=numero_acto,
            lugar_entrega=lugar_entrega,
            tiempo_entrega=tiempo_entrega,
        )
        # Ajuste específico de Doc_Gen:
        # medidas/no-incapacidad/desglose/nota/carta usan "representante_documentos".
        if any(
            token in output_name
            for token in (
                "medidas_de_retorsion",
                "no_incapacidad_para_contratar",
                "desglose_de_precios",
                "nota_adicional",
                "carta_de_adhesion",
            )
        ):
            replacements["[Representante_legal_de_la_Entidad_Licitante]"] = str(
                representante_legal_documentos or ""
            ).strip()
        template_path = _resolve_doc_gen_template(template_name)
        file_bytes = _fill_docx_placeholders(template_path, replacements)
        docs_out.append(
            {
                "template": template_name,
                "file_name": output_name,
                "bytes": file_bytes,
            }
        )
    return docs_out


def _get_lp_doc_folder(drive, empresa_full: str) -> str:
    _, company_folders = _get_drive_folders(drive)
    company_folder = company_folders.get(empresa_full)
    if not company_folder:
        raise RuntimeError(f"No se encontró carpeta de Drive para {empresa_full}.")
    return _find_or_create_folder(drive, LP_DOC_FOLDER_NAME, company_folder)


def _upload_quote_html(
    drive,
    folder_id: str,
    filename: str,
    html_body: str,
    existing_file_id: str | None = None,
) -> dict:
    media = MediaIoBaseUpload(BytesIO(html_body.encode("utf-8")), mimetype="text/html", resumable=False)
    if existing_file_id:
        return drive.files().update(
            fileId=existing_file_id,
            media_body=media,
            fields="id,name",
            supportsAllDrives=True,
        ).execute()
    metadata = {"name": filename, "parents": [folder_id]}
    return drive.files().create(
        body=metadata,
        media_body=media,
        fields="id,name",
        supportsAllDrives=True,
    ).execute()


def _download_drive_file(drive, file_id: str) -> bytes:
    request = drive.files().get_media(fileId=file_id, supportsAllDrives=True)
    fh = BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    return fh.getvalue()


def _items_resumen(items_df: pd.DataFrame) -> str:
    if items_df.empty:
        return ""
    first = str(items_df.iloc[0].get("producto_servicio", "") or "").strip()
    restantes = max(len(items_df) - 1, 0)
    if restantes:
        return f"{first} (+{restantes} más)"
    return first
def _build_items_dataframe(raw: pd.DataFrame) -> pd.DataFrame:
    df = raw.copy()
    if "cantidad" in df.columns:
        df["cantidad"] = pd.to_numeric(df["cantidad"], errors="coerce").fillna(0.0)
    if "precio_unitario" in df.columns:
        df["precio_unitario"] = pd.to_numeric(df["precio_unitario"], errors="coerce").fillna(0.0)
    df["importe"] = df["cantidad"] * df["precio_unitario"]
    return df


def _guess_mime_from_filename(filename: str) -> str:
    ext = Path(filename or "").suffix.lower()
    if ext == ".xlsx":
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if ext == ".html":
        return "text/html"
    if ext == ".pdf":
        return "application/pdf"
    return "application/octet-stream"


def _upload_drive_binary(
    drive,
    folder_id: str,
    filename: str,
    data: bytes,
    mime_type: str,
    existing_file_id: str | None = None,
) -> dict:
    media = MediaIoBaseUpload(BytesIO(data), mimetype=mime_type, resumable=False)
    if existing_file_id:
        return drive.files().update(
            fileId=existing_file_id,
            media_body=media,
            fields="id,name",
            supportsAllDrives=True,
        ).execute()
    metadata = {"name": filename, "parents": [folder_id]}
    return drive.files().create(
        body=metadata,
        media_body=media,
        fields="id,name",
        supportsAllDrives=True,
    ).execute()


def _summarize_quote_title(items: pd.DataFrame, details: str) -> str:
    base = ""
    if not items.empty:
        base = str(items.iloc[0].get("producto_servicio", "") or "").strip()
    if not base:
        base = str(details or "").strip()
    base = " ".join(base.split())
    if not base:
        return "Cotizacion de bienes y servicios"
    return base[:55].rstrip()


def _build_standard_quote_excel(
    empresa: str,
    numero_cot: str,
    fecha_cot: date,
    cliente: str,
    direccion: str,
    cliente_ruc: str,
    cliente_dv: str,
    items_df: pd.DataFrame,
    impuesto_pct: float,
    condiciones: Dict[str, str],
    detalles_extra: str,
    numero_excel: Optional[str] = None,
    titulo_override: Optional[str] = None,
    allow_blank_ruc_dv: bool = False,
) -> bytes:
    if empresa == "RIR Medical":
        template_path = TEMPLATE_RIR_STANDARD
        header_path = HEADER_RIR_STANDARD
    else:
        template_path = TEMPLATE_RS_STANDARD
        header_path = HEADER_RS_STANDARD

    if not template_path.exists():
        raise FileNotFoundError(f"No se encontró la plantilla: {template_path}")

    wb = load_workbook(template_path)
    ws = wb["cotizacion"] if "cotizacion" in wb.sheetnames else wb[wb.sheetnames[0]]

    items = items_df.copy()
    if "producto_servicio" not in items.columns:
        items["producto_servicio"] = ""
    items["producto_servicio"] = items["producto_servicio"].fillna("").astype(str).str.strip()
    if "unidad" not in items.columns:
        items["unidad"] = "UND"
    items["unidad"] = items["unidad"].fillna("UND").astype(str).str.strip().replace("", "UND")
    items["cantidad"] = pd.to_numeric(items.get("cantidad"), errors="coerce").fillna(0.0)
    items["precio_unitario"] = pd.to_numeric(items.get("precio_unitario"), errors="coerce").fillna(0.0)
    items = items[
        (items["producto_servicio"].str.len() > 0)
        | (items["cantidad"] > 0)
        | (items["precio_unitario"] > 0)
    ].reset_index(drop=True)
    if items.empty:
        raise ValueError("Debes agregar al menos un item para generar el Excel.")

    numero_items = len(items)
    filas_a_insertar = max(numero_items - 1, 0)
    fila_inicio_items = 23
    if filas_a_insertar:
        ws.insert_rows(fila_inicio_items + 1, filas_a_insertar)

    borde_sencillo = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_right = Alignment(horizontal="right", vertical="center", wrap_text=True)

    subtotal = 0.0
    for idx, row in items.iterrows():
        excel_row = fila_inicio_items + idx
        cantidad = float(row.get("cantidad", 0.0) or 0.0)
        precio_unitario = float(row.get("precio_unitario", 0.0) or 0.0)
        total_item = round(cantidad * precio_unitario, 2)
        subtotal = round(subtotal + total_item, 2)

        ws[f"B{excel_row}"] = idx + 1
        ws[f"C{excel_row}"] = str(row.get("producto_servicio", "") or "")
        ws[f"D{excel_row}"] = str(row.get("unidad", "UND") or "UND")
        ws[f"E{excel_row}"] = cantidad
        ws[f"F{excel_row}"] = precio_unitario
        # Mantener fórmula editable en Excel para que recalcule al editar cantidad/costo.
        ws[f"G{excel_row}"] = f"=IFERROR(E{excel_row}*F{excel_row},0)"

        ws[f"B{excel_row}"].alignment = align_center
        ws[f"C{excel_row}"].alignment = Alignment(wrap_text=True, vertical="center")
        ws[f"D{excel_row}"].alignment = align_center
        ws[f"E{excel_row}"].alignment = align_center
        ws[f"F{excel_row}"].alignment = align_right
        ws[f"G{excel_row}"].alignment = align_right
        ws[f"F{excel_row}"].number_format = "$#,##0.00"
        ws[f"G{excel_row}"].number_format = "$#,##0.00"

        for col in ("B", "C", "D", "E", "F", "G"):
            ws[f"{col}{excel_row}"].border = borde_sencillo

    fila_subtotal = fila_inicio_items + numero_items
    fila_impuesto = fila_subtotal + 1
    fila_total = fila_subtotal + 2
    last_item_row = fila_inicio_items + numero_items - 1
    impuesto_factor = float(impuesto_pct) / 100.0

    ws[f"F{fila_subtotal}"] = "Subtotal ="
    ws[f"F{fila_impuesto}"] = f"Impuesto ({impuesto_pct:.2f}%) ="
    ws[f"F{fila_total}"] = "Total ="
    # Totales con fórmula para mantener trazabilidad y edición manual.
    ws[f"G{fila_subtotal}"] = f"=SUM(G{fila_inicio_items}:G{last_item_row})"
    ws[f"G{fila_impuesto}"] = f"=G{fila_subtotal}*{impuesto_factor:.6f}"
    ws[f"G{fila_total}"] = f"=G{fila_subtotal}+G{fila_impuesto}"

    for row_tot in (fila_subtotal, fila_impuesto, fila_total):
        ws[f"F{row_tot}"].alignment = align_right
        ws[f"G{row_tot}"].alignment = align_right
        ws[f"F{row_tot}"].border = borde_sencillo
        ws[f"G{row_tot}"].border = borde_sencillo
        ws[f"G{row_tot}"].number_format = "$#,##0.00"

    title = str(titulo_override or "").strip() or _summarize_quote_title(items, detalles_extra)
    numero_visible = str(numero_excel or numero_cot or "").strip() or str(numero_cot)
    ws["B13"] = cliente or "-"
    ws["G13"] = fecha_cot.strftime("%Y-%m-%d")
    if allow_blank_ruc_dv:
        ruc_text = str(cliente_ruc or "").strip()
        dv_text = str(cliente_dv or "").strip()
    else:
        ruc_text = str(cliente_ruc or "-").strip()
        dv_text = str(cliente_dv or "-").strip()
    ws["B14"] = f"RUC: {ruc_text}   DV: {dv_text}"
    ws["E18"] = numero_visible
    ws["C19"] = title
    ws["B21"] = title
    if numero_visible != str(numero_cot or "").strip():
        ws["B12"] = str(numero_cot or "").strip()

    forma_pago = condiciones.get("Condicion de pago") or "Credito"
    entrega = condiciones.get("Entrega") or "15 días hábiles"
    lugar_entrega = condiciones.get("Lugar de entrega") or "-"
    vigencia = condiciones.get("Vigencia") or "15 días"

    fila_lugar = 30 + numero_items
    ws[f"B{fila_lugar - 1}"] = f"Forma de pago: {forma_pago}"
    ws[f"B{fila_lugar}"] = f"Lugar de entrega: {lugar_entrega}"
    ws[f"B{fila_lugar + 1}"] = f"Tiempo de entrega: {entrega}"
    ws[f"B{fila_lugar + 2}"] = "Garantía: De fábrica"
    ws[f"B{fila_lugar + 3}"] = "Adjudicación: Global"
    ws[f"B{fila_lugar + 4}"] = f"Validez de la propuesta: {vigencia}"
    if direccion:
        ws[f"B{fila_lugar + 5}"] = f"Dirección del cliente: {direccion}"

    extra_lines = [line.strip() for line in str(detalles_extra or "").splitlines() if line.strip()]
    extra_row = fila_lugar + 6
    if extra_lines:
        ws[f"B{extra_row}"] = "Observaciones:"
        for idx, line in enumerate(extra_lines[:8], start=1):
            ws[f"B{extra_row + idx}"] = line
        base_firma_row = extra_row + min(len(extra_lines), 8) + 3
    else:
        base_firma_row = fila_lugar + 9

    firma_row = max(38 + filas_a_insertar, base_firma_row)
    ws[f"B{firma_row - 2}"] = "Atentamente,"
    ws[f"B{firma_row + 5}"] = "ING RODRIGO SÁNCHEZ"
    ws[f"B{firma_row + 6}"] = "Representante Legal"

    if header_path.exists():
        encabezado = Image(str(header_path))
        encabezado.width, encabezado.height = 590, 202
        ws.add_image(encabezado, "B2")

    if SIGNATURE_STANDARD.exists():
        firma = Image(str(SIGNATURE_STANDARD))
        firma.width, firma.height = 150, 100
        ws.add_image(firma, f"B{firma_row}")

    # Impresión amigable: ancho a 1 página y alto automático para tablas largas.
    ws.page_setup.orientation = "portrait"
    ws.page_setup.paperSize = 9  # A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.page_margins.left = 0.3
    ws.page_margins.right = 0.3
    ws.page_margins.top = 0.3
    ws.page_margins.bottom = 0.3
    ws.page_margins.header = 0.2
    ws.page_margins.footer = 0.2
    ws.print_title_rows = "1:22"

    output = BytesIO()
    wb.save(output)
    wb.close()
    return output.getvalue()


def _extract_standard_excel_preview(excel_bytes: bytes) -> dict:
    def _is_formula(value) -> bool:
        return isinstance(value, str) and value.strip().startswith("=")

    def _to_float(value) -> float:
        if value is None:
            return 0.0
        if isinstance(value, (int, float)):
            return float(value)
        text = str(value).strip()
        if not text:
            return 0.0
        if text.startswith("="):
            return 0.0
        text = (
            text.replace("B/.", "")
            .replace("B/", "")
            .replace("$", "")
            .replace("USD", "")
            .strip()
        )
        if "," in text and "." in text:
            text = text.replace(",", "")
        elif "," in text and "." not in text:
            text = text.replace(",", ".")
        text = re.sub(r"[^0-9.\-]", "", text)
        if text in {"", "-", ".", "-."}:
            return 0.0
        try:
            return float(text)
        except (TypeError, ValueError):
            return 0.0

    wb = load_workbook(BytesIO(excel_bytes), data_only=False)
    ws = wb["cotizacion"] if "cotizacion" in wb.sheetnames else wb[wb.sheetnames[0]]

    items = []
    row = 23
    max_rows = 800
    subtotal_row = None
    while row < max_rows:
        item = ws[f"B{row}"].value
        desc = ws[f"C{row}"].value
        unidad = ws[f"D{row}"].value
        cantidad = ws[f"E{row}"].value
        precio_unitario = ws[f"F{row}"].value
        total = ws[f"G{row}"].value
        marker = str(precio_unitario or "").strip().lower()
        if any(token in marker for token in ("subtotal", "impuesto", "total")):
            subtotal_row = row
            break
        if not any([item, desc, unidad, cantidad, precio_unitario, total]):
            break
        cantidad_val = _to_float(cantidad)
        unit_val = _to_float(precio_unitario)
        total_val = round(cantidad_val * unit_val, 2) if _is_formula(total) else _to_float(total)

        items.append(
            {
                "Item": item,
                "Descripción": desc or "",
                "Unidad": unidad or "",
                "Cantidad": cantidad_val,
                "Costo Unitario": unit_val,
                "Total": total_val,
            }
        )
        row += 1

    if subtotal_row is None:
        subtotal_row = row
        # Fallback: buscar explícitamente la fila de Subtotal.
        for r in range(23, max_rows):
            marker = str(ws[f"F{r}"].value or "").strip().lower()
            if "subtotal" in marker:
                subtotal_row = r
                break

    subtotal_cell = ws[f"G{subtotal_row}"].value
    impuesto_cell = ws[f"G{subtotal_row + 1}"].value
    total_cell = ws[f"G{subtotal_row + 2}"].value

    subtotal_calc = round(sum(float(x.get("Total", 0.0) or 0.0) for x in items), 2)
    subtotal = subtotal_calc if _is_formula(subtotal_cell) else _to_float(subtotal_cell)

    impuesto_label = str(ws[f"F{subtotal_row + 1}"].value or "")
    pct_match = re.search(r"([0-9]+(?:[.,][0-9]+)?)\s*%", impuesto_label)
    impuesto_pct = 0.0
    if pct_match:
        try:
            impuesto_pct = float(str(pct_match.group(1)).replace(",", "."))
        except ValueError:
            impuesto_pct = 0.0
    impuesto_calc = round(subtotal * (impuesto_pct / 100.0), 2)
    impuesto = impuesto_calc if _is_formula(impuesto_cell) else _to_float(impuesto_cell)
    total_calc = round(subtotal + impuesto, 2)
    total_doc = total_calc if _is_formula(total_cell) else _to_float(total_cell)

    conditions = []
    for r in range(28, 90):
        text = str(ws[f"B{r}"].value or "").strip()
        if text and ":" in text:
            conditions.append(text)

    out = {
        "numero": str(ws["E18"].value or ""),
        "titulo": str(ws["C19"].value or ""),
        "cliente": str(ws["B13"].value or ""),
        "fecha": str(ws["G13"].value or ""),
        "ruc_dv": str(ws["B14"].value or ""),
        "items": items,
        "subtotal": subtotal,
        "impuesto": impuesto,
        "total": total_doc,
        "condiciones": conditions,
    }
    wb.close()
    return out


def _save_panama_quote_to_history(
    *,
    client,
    creds,
    sheet_id: str,
    cot_df: pd.DataFrame,
    manual_request_id: str,
    empresa_short: str,
    enlace_pc: str,
    titulo_excel: str,
    panama_meta: dict[str, Any],
    items_panama_df: pd.DataFrame,
    paga_itbms: bool,
    presupuesto_df: pd.DataFrame,
    costo_interno: float,
    factor_ganancia: float,
    precio_cotizar: float,
    ganancia: float,
    financiamiento_tipo: str,
    financiamiento_interes_pct: float,
    costo_financiamiento: float,
    ganancia_neta: float,
    tiempo_inversion: float,
    inversion_etapa_1: float,
    tiempo_intermedio: float,
    inversion_etapa_intermedia: float,
    tiempo_cobro: float,
    inversion_etapa_2: float,
    manual_cliente_ruc: str = "",
    manual_cliente_dv: str = "",
) -> dict[str, Any]:
    _ensure_cotizaciones_sheet(client, sheet_id)
    df_write = _normalize_cotizaciones_df(read_worksheet(client, sheet_id, SHEET_NAME_COT))
    if df_write.empty and not cot_df.empty:
        df_write = cot_df.copy()

    row_id = f"pc_manual_{manual_request_id}"
    now = datetime.now().isoformat(timespec="seconds")
    existing_row = None
    if not df_write.empty and row_id in df_write["id"].astype(str).values:
        existing_row = df_write[df_write["id"].astype(str) == row_id].iloc[0].to_dict()

    empresa = _company_full_from_short(empresa_short)
    prefijo = existing_row.get("prefijo") if existing_row else ""
    seq = 0
    numero_cot = ""
    if existing_row:
        prefijo = str(existing_row.get("prefijo") or _pc_prefijo_from_short(empresa_short))
        try:
            seq = int(float(existing_row.get("secuencia") or 0))
        except (TypeError, ValueError):
            seq = 0
        numero_cot = str(existing_row.get("numero_cotizacion") or "")
    if not numero_cot:
        prefijo = _pc_prefijo_from_short(empresa_short)
        seq = _next_sequence(df_write, prefijo)
        numero_cot = _build_numero_cot(prefijo, seq)

    items_df = _build_standard_items_from_panama(items_panama_df)
    if items_df.empty:
        raise ValueError("La cotización de Panamá Compra no trajo ítems válidos.")

    impuesto_pct = 7.0 if paga_itbms else 0.0
    numero_acto = str(panama_meta.get("numero_acto") or "").strip()
    titulo_resumido = str(titulo_excel or panama_meta.get("titulo") or "").strip()
    fecha_cot = _parse_panama_fecha(panama_meta.get("fecha"))
    cliente_nombre = (
        str(panama_meta.get("entidad") or "").strip()
        or (str(titulo_excel or "").strip() or "Cliente Panamá Compra")
    )[:120]
    direccion = ""
    cliente_ruc = str(manual_cliente_ruc or "").strip() or str(panama_meta.get("ruc") or "").strip()
    cliente_dv = str(manual_cliente_dv or "").strip() or str(panama_meta.get("dv") or "").strip()

    condiciones = {
        "Vigencia": str(panama_meta.get("vigencia") or "120 días"),
        "Condicion de pago": str(panama_meta.get("forma_pago") or "Credito"),
        "Entrega": str(panama_meta.get("tiempo_entrega") or "Según pliego"),
        "Lugar de entrega": str(panama_meta.get("lugar_entrega") or "Según Panamá Compra"),
    }
    # En Panamá Compra no mostramos bloque de observaciones automáticas en la plantilla final.
    detalles_extra = ""

    excel_bytes = _build_standard_quote_excel(
        empresa=empresa,
        numero_cot=numero_cot,
        fecha_cot=fecha_cot,
        cliente=cliente_nombre,
        direccion=direccion,
        cliente_ruc=cliente_ruc,
        cliente_dv=cliente_dv,
        items_df=items_df,
        impuesto_pct=impuesto_pct,
        condiciones=condiciones,
        detalles_extra=detalles_extra,
        numero_excel=numero_acto or numero_cot,
        titulo_override=titulo_resumido or None,
        allow_blank_ruc_dv=True,
    )

    subtotal = float((items_df["cantidad"] * items_df["precio_unitario"]).sum())
    impuesto_monto = subtotal * (impuesto_pct / 100.0)
    total = subtotal + impuesto_monto
    tiempo_recuperacion = tiempo_inversion + tiempo_intermedio + tiempo_cobro

    items_json = json.dumps(
        items_df[["producto_servicio", "unidad", "cantidad", "precio_unitario"]].to_dict(orient="records"),
        ensure_ascii=False,
    )
    presupuesto_items_json = json.dumps(
        presupuesto_df[["producto_servicio", "cantidad", "precio_unitario"]].to_dict(orient="records"),
        ensure_ascii=False,
    )
    condiciones_json = json.dumps(condiciones, ensure_ascii=False)

    descripcion_corta = _generate_quote_short_description(
        tipo_cotizacion="Panama Compra",
        empresa=empresa,
        cliente=cliente_nombre,
        detalles=detalles_extra,
        items=[str(x).strip() for x in items_df["producto_servicio"].tolist() if str(x).strip()][:5],
    )
    excel_filename = _quote_excel_filename(numero_cot, descripcion_corta)

    drive_file_id = existing_row.get("drive_file_id") if existing_row else ""
    drive_file_name = existing_row.get("drive_file_name") if existing_row else ""
    drive_file_url = existing_row.get("drive_file_url") if existing_row else ""
    drive_folder = existing_row.get("drive_folder") if existing_row else ""
    presupuesto_drive_file_id = existing_row.get("presupuesto_drive_file_id") if existing_row else ""
    presupuesto_drive_file_name = existing_row.get("presupuesto_drive_file_name") if existing_row else ""
    presupuesto_drive_file_url = existing_row.get("presupuesto_drive_file_url") if existing_row else ""

    if creds is not None:
        drive = _get_drive_client(creds)
        _, folders = _get_drive_folders(drive)
        folder_id = folders.get(empresa)
        if folder_id:
            upload = _upload_drive_binary(
                drive,
                folder_id,
                excel_filename,
                excel_bytes,
                _guess_mime_from_filename(excel_filename),
                existing_file_id=drive_file_id or None,
            )
            drive_file_id = upload.get("id", drive_file_id)
            drive_file_name = upload.get("name", excel_filename)
            drive_folder = folder_id
            if drive_file_id:
                drive_file_url = f"https://drive.google.com/file/d/{drive_file_id}/view"

            presupuesto_html = _build_budget_html(
                empresa=empresa,
                numero=numero_cot,
                fecha_cot=fecha_cot,
                presupuesto_df=presupuesto_df,
                costo_interno=costo_interno,
                factor_ganancia=factor_ganancia,
                precio_cotizar=precio_cotizar,
                ganancia=ganancia,
                financiamiento_tipo=financiamiento_tipo,
                financiamiento_interes_pct=financiamiento_interes_pct,
                costo_financiamiento=costo_financiamiento,
                ganancia_neta=ganancia_neta,
                tiempo_inversion=tiempo_inversion,
                inversion_etapa_1=inversion_etapa_1,
                tiempo_inicio_ejecucion_presentacion=tiempo_intermedio,
                inversion_etapa_intermedia=inversion_etapa_intermedia,
                tiempo_cobro=tiempo_cobro,
                inversion_etapa_2=inversion_etapa_2,
            )
            presupuesto_filename = f"Presupuesto_{numero_cot}.html"
            presupuesto_upload = _upload_quote_html(
                drive,
                folder_id,
                presupuesto_filename,
                presupuesto_html,
                existing_file_id=presupuesto_drive_file_id or None,
            )
            presupuesto_drive_file_id = presupuesto_upload.get("id", presupuesto_drive_file_id)
            presupuesto_drive_file_name = presupuesto_upload.get("name", presupuesto_filename)
            if presupuesto_drive_file_id:
                presupuesto_drive_file_url = f"https://drive.google.com/file/d/{presupuesto_drive_file_id}/view"

    row = {
        "id": row_id,
        "numero_cotizacion": numero_cot,
        "prefijo": prefijo,
        "secuencia": seq,
        "empresa": empresa,
        "tipo_cotizacion": "Panama Compra",
        "descripcion_corta": descripcion_corta,
        "cliente_nombre": cliente_nombre,
        "cliente_direccion": direccion,
        "cliente_ruc": cliente_ruc,
        "cliente_dv": cliente_dv,
        "fecha_cotizacion": fecha_cot.isoformat(),
        "created_at": existing_row.get("created_at") if existing_row else now,
        "updated_at": now,
        "moneda": "USD",
        "subtotal": subtotal,
        "impuesto_pct": impuesto_pct,
        "impuesto_monto": impuesto_monto,
        "total": total,
        "items_json": items_json,
        "items_resumen": _items_resumen(items_df),
        "detalles_extra": detalles_extra,
        "presupuesto_items_json": presupuesto_items_json,
        "presupuesto_subtotal": costo_interno,
        "presupuesto_factor_ganancia": factor_ganancia,
        "presupuesto_precio_cotizar": precio_cotizar,
        "presupuesto_ganancia": ganancia,
        "presupuesto_financiamiento_tipo": financiamiento_tipo,
        "presupuesto_financiamiento_interes_pct": financiamiento_interes_pct,
        "presupuesto_costo_financiamiento": costo_financiamiento,
        "presupuesto_ganancia_neta": ganancia_neta,
        "presupuesto_t_inversion_presentacion": tiempo_inversion,
        "presupuesto_inversion_etapa_1": inversion_etapa_1,
        "presupuesto_t_inicio_ejecucion_presentacion": tiempo_intermedio,
        "presupuesto_inversion_etapa_intermedia": inversion_etapa_intermedia,
        "presupuesto_t_presentacion_cobro": tiempo_cobro,
        "presupuesto_inversion_etapa_2": inversion_etapa_2,
        "presupuesto_t_recuperacion": tiempo_recuperacion,
        "condiciones_json": condiciones_json,
        "vigencia": condiciones["Vigencia"],
        "forma_pago": condiciones["Condicion de pago"],
        "entrega": condiciones["Entrega"],
        "lugar_entrega": condiciones["Lugar de entrega"],
        "estado": existing_row.get("estado", "vigente") if existing_row else "vigente",
        "notas": existing_row.get("notas", "") if existing_row else "",
        "drive_file_id": drive_file_id,
        "drive_file_name": drive_file_name,
        "drive_file_url": drive_file_url,
        "drive_folder": drive_folder,
        "presupuesto_drive_file_id": presupuesto_drive_file_id,
        "presupuesto_drive_file_name": presupuesto_drive_file_name,
        "presupuesto_drive_file_url": presupuesto_drive_file_url,
    }

    if existing_row and row_id in df_write["id"].astype(str).values:
        idx = df_write.index[df_write["id"].astype(str) == row_id][0]
        for col in COT_COLUMNS:
            df_write.at[idx, col] = row.get(col, "")
    else:
        df_write = pd.concat([df_write, pd.DataFrame([row])], ignore_index=True)
    write_worksheet(client, sheet_id, SHEET_NAME_COT, _normalize_cotizaciones_df(df_write))

    return {
        "row": row,
        "numero_cotizacion": numero_cot,
        "excel_bytes": excel_bytes,
        "excel_name": excel_filename,
    }


def _build_invoice_html(
    empresa: str,
    branding: Dict[str, str],
    numero: str,
    fecha_cot: date,
    cliente: str,
    direccion: str,
    cliente_ruc: str,
    cliente_dv: str,
    firma_b64: str,
    detalles_extra: str,
    layout_extra_space: int,
    layout_spacers: Dict[str, int],
    items: pd.DataFrame,
    impuesto_pct: float,
    condiciones: Dict[str, str],
) -> str:
    logo_b64 = branding.get("logo_b64", "")
    background_b64 = branding.get("background_b64", "")
    contacto_html = branding.get("contacto_html", "")
    logo_scale = float(branding.get("logo_scale", 1.0))
    logo_box_width = int(branding.get("logo_box_width", branding.get("logo_box", 190)) * logo_scale)
    logo_box_height = int(branding.get("logo_box_height", branding.get("logo_box", 190)) * logo_scale)
    logo_width = int(branding.get("logo_width", branding.get("logo_size", 180)) * logo_scale)
    logo_height = int(branding.get("logo_height", branding.get("logo_size", 180)) * logo_scale)
    logo_left = int(branding.get("logo_left", 120))
    logo_top = int(branding.get("logo_top", 120))
    header_left = int(branding.get("header_left", logo_left + logo_box_width + 30))
    header_top = int(branding.get("header_top", 140))
    header_width = int(branding.get("header_width", 520))
    header_height = int(branding.get("header_height", logo_box_height))
    content_offset_x = int(branding.get("content_offset_x", 0))
    content_offset_y = int(branding.get("content_offset_y", 0))
    layout_spacers = layout_spacers or {}
    layout_global_offset = int(layout_spacers.get("global_offset", 0))
    title_offset = int(layout_spacers.get("title_offset", 0))
    space_after_title = int(layout_spacers.get("space_after_title", 0))
    space_after_columns = int(layout_spacers.get("space_after_columns", 0))
    space_after_table = int(layout_spacers.get("space_after_table", 0))
    space_after_totals = int(layout_spacers.get("space_after_totals", 0))
    space_after_extra = int(layout_spacers.get("space_after_extra", 0))
    space_after_conditions = int(layout_spacers.get("space_after_conditions", 0))

    title_top = 380 + content_offset_y + layout_global_offset + title_offset
    title_left = 120 + content_offset_x
    title_meta_top = 440 + content_offset_y + layout_global_offset + title_offset
    title_meta_left = 120 + content_offset_x
    columns_top = 520 + content_offset_y + layout_global_offset + space_after_title
    columns_left = 120 + content_offset_x
    table_top = 720 + content_offset_y + layout_global_offset + space_after_title + space_after_columns
    table_left = 120 + content_offset_x
    totals_right = 160 - content_offset_x
    conditions_left = 120 + content_offset_x
    extra_left = 120 + content_offset_x

    subtotal = float(items["importe"].sum())
    impuesto = subtotal * (impuesto_pct / 100.0)
    total = subtotal + impuesto
    cliente_ruc_text = ""
    if (cliente_ruc or "").strip() or (cliente_dv or "").strip():
        cliente_ruc_text = f"RUC: {cliente_ruc or '-'} DV: {cliente_dv or '-'}"
    else:
        cliente_ruc_text = "-"

    rows: List[str] = []
    row_height_base = 44
    line_height = 18
    table_rows_height = 0
    for _, row in items.iterrows():
        producto_text = str(row.get("producto_servicio", "") or "")
        line_count = max(1, math.ceil(len(producto_text) / 60))
        table_rows_height += row_height_base + (line_count - 1) * line_height
        rows.append(
            f"""
            <tr>
              <td>{html.escape(producto_text)}</td>
              <td class=\"num\">{row.get('cantidad', 0):,.0f}</td>
              <td class=\"num\">{_format_money(row.get('precio_unitario', 0))}</td>
              <td class=\"num\">{_format_money(row.get('importe', 0))}</td>
            </tr>
            """
        )

    if not rows:
        table_rows_height = row_height_base

    table_height = 46 + table_rows_height
    totals_top = table_top + table_height + 40 + space_after_table

    extra_text = (detalles_extra or "").strip()
    extra_lines = 0
    if extra_text:
        for line in extra_text.splitlines() or [""]:
            extra_lines += max(1, math.ceil(len(line) / 90))
    extra_height = extra_lines * 20 + 30 if extra_text else 0
    extra_top = totals_top + 120 + space_after_totals
    conditions_top = (extra_top + extra_height + 30 if extra_text else totals_top + 120) + layout_extra_space + space_after_extra
    conditions_lines = 0
    for label, text in condiciones.items():
        combined = f"{label}: {text}"
        conditions_lines += max(1, math.ceil(len(combined) / 90))
    conditions_height = 40 + conditions_lines * 20
    signature_height = 200
    signature_top = conditions_top + conditions_height + 30 + space_after_conditions

    base_page_height = 2000
    header_clearance = max(logo_top + logo_box_height, header_top + header_height) + 40
    content_top_offset = header_clearance
    bottom_margin = 360
    block_bottom = signature_top + signature_height
    page_index = int(block_bottom // base_page_height)
    page_limit = (page_index + 1) * base_page_height
    if block_bottom + bottom_margin > page_limit:
        conditions_top = (page_index + 1) * base_page_height
        signature_top = conditions_top + conditions_height + 30 + space_after_conditions
        block_bottom = signature_top + signature_height
    content_bottom = block_bottom + bottom_margin
    page_count = max(1, math.ceil(content_bottom / base_page_height))
    page_height = page_count * base_page_height

    def _apply_page_offset(value: float) -> float:
        page_idx = int(value // base_page_height)
        if page_idx <= 0:
            return value
        return value + page_idx * content_top_offset

    title_top = _apply_page_offset(title_top)
    title_meta_top = _apply_page_offset(title_meta_top)
    columns_top = _apply_page_offset(columns_top)
    table_top = _apply_page_offset(table_top)
    totals_top = _apply_page_offset(totals_top)
    extra_top = _apply_page_offset(extra_top)
    conditions_top = _apply_page_offset(conditions_top)
    signature_top = _apply_page_offset(signature_top)
    signature_img = ""
    if firma_b64:
        signature_img = (
            "<img src='data:image/png;base64," + firma_b64 + "' alt='firma' />"
        )
    signature_html = (
        "<div class=\"signature\" style=\"top:"
        + str(signature_top)
        + "px;left:"
        + str(conditions_left)
        + "px;\">"
        + signature_img
        + "<div class=\"signature-name\">Rodrigo S&aacute;nchez P.</div>"
        + "<div class=\"signature-id\">C&eacute;dula: 9-740-624</div>"
        + "</div>"
    )
    header_repeats = ""
    if page_count > 1:
        for page in range(1, page_count):
            offset = page * base_page_height
            header_repeats += (
                "  <div class=\"logo page-header\" style=\"left:"
                + str(logo_left)
                + "px;top:"
                + str(logo_top + offset)
                + "px;width:"
                + str(logo_box_width)
                + "px;height:"
                + str(logo_box_height)
                + "px;\">\n"
                + ("    <img src='data:image/png;base64,"
                   + logo_b64
                   + "' alt='logo' style='width:"
                   + str(logo_width)
                   + "px;height:"
                   + str(logo_height)
                   + "px;' />\n" if logo_b64 else "")
                + "  </div>\n"
                + "  <div class=\"header-info page-header\" style=\"left:"
                + str(header_left)
                + "px;top:"
                + str(header_top + offset)
                + "px;width:"
                + str(header_width)
                + "px;height:"
                + str(header_height)
                + "px;\">\n"
                + "    <div class=\"empresa\">" + html.escape(empresa) + "</div>\n"
                + "    <div class=\"datos\">" + contacto_html + "</div>\n"
                + "  </div>\n"
            )

    sample_rows = "".join(rows) or """
        <tr>
            <td colspan=\"4\" style=\"text-align:center;color:#64748b;\">Agrega items para ver el desglose.</td>
        </tr>
    """

    condiciones_html = "".join(
        f"<li><strong>{html.escape(label)}:</strong> {html.escape(text)}</li>"
        for label, text in condiciones.items()
    )
    extra_html = ""
    if extra_text:
        extra_html = (
            "<div class=\"extra-details\" style=\"top:"
            + str(extra_top)
            + "px;left:"
            + str(extra_left)
            + "px;\">"
            + "<h4>Detalles adicionales</h4>"
            + "<div class=\"extra-body\">"
            + html.escape(extra_text).replace("\n", "<br>")
            + "</div></div>"
        )
    return f"""
<style>
  @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800&family=Manrope:wght@400;600;700;800&display=swap');
  .quote-page {{
    position: relative;
    width: 1414px;
    height: {page_height}px;
    min-height: 2000px;
    margin: 0 auto 24px auto;
    background: #ffffff url('data:image/png;base64,{background_b64}') top center / 100% 2000px repeat-y;
    font-family: 'Manrope', 'Inter', 'Segoe UI', sans-serif;
    color: #0c2349;
    -webkit-print-color-adjust: exact;
    print-color-adjust: exact;
  }}
  .logo {{
    position: absolute;
    top: 120px;
    left: 120px;
    width: 190px;
    height: 190px;
    display: flex;
    align-items: center;
    justify-content: center;
  }}
  .logo img {{
    width: 180px;
    height: 180px;
    object-fit: contain;
  }}
  .header-info {{
    position: absolute;
    top: 140px;
    left: 340px;
    width: 520px;
    color: #6b7280;
    line-height: 1.35;
    display: flex;
    flex-direction: column;
    justify-content: center;
  }}
  .header-info .empresa {{
    font-size: 28px;
    font-weight: 800;
    color: #4b5563;
    margin: 0 0 8px 0;
  }}
  .header-info .datos {{
    font-size: 16px;
    color: #6b7280;
  }}
  .title {{
    position: absolute;
    top: 380px;
    left: 120px;
    font-size: 40px;
    font-weight: 800;
  }}
  .title-meta {{
    position: absolute;
    top: 440px;
    left: 120px;
    font-size: 16px;
    color: #6b7280;
    line-height: 1.4;
  }}
  .columns {{
    position: absolute;
    top: 520px;
    left: 120px;
    right: 120px;
    display: grid;
    grid-template-columns: 1fr 1fr;
    gap: 70px;
    font-size: 16px;
    line-height: 1.45;
  }}
  .columns h4 {{
    margin: 0 0 10px 0;
    font-size: 17px;
    color: #0c2349;
  }}
  .columns .block div {{
    margin: 0 0 6px 0;
    color: #1f2f46;
  }}
  .columns .emisor-name {{
    margin-bottom: 6px;
  }}
  .columns .contacto {{
    color: #1f2f46;
  }}
  .table-wrap {{
    position: absolute;
    top: 720px;
    left: 120px;
    width: 1174px;
  }}
  .extra-details {{
    position: absolute;
    width: 1174px;
    font-size: 15px;
    line-height: 1.5;
    color: #1f2f46;
  }}
  .extra-details h4 {{
    margin: 0 0 10px 0;
    font-size: 16px;
    font-weight: 800;
    color: #0c2349;
  }}
  table.items {{
    width: 100%;
    border-collapse: collapse;
    font-size: 15px;
    color: #0c2349;
  }}
  table.items th {{
    background: #1c336a;
    color: #ffffff;
    padding: 12px 10px;
    text-align: left;
    border: 2px solid #1c336a;
    font-weight: 700;
  }}
  table.items td {{
    border: 1px solid #d9e2f1;
    padding: 12px 10px;
    vertical-align: top;
  }}
  table.items td.num {{
    text-align: center;
    white-space: nowrap;
  }}
  .totals {{
    position: absolute;
    top: 1180px;
    right: 160px;
    width: 320px;
    font-size: 16px;
    line-height: 1.6;
  }}
  .totals div {{
    display: flex;
    justify-content: space-between;
  }}
  .totals .total {{
    font-weight: 800;
    font-size: 20px;
  }}
  .conditions {{
    position: absolute;
    top: 1340px;
    left: 120px;
    width: 1174px;
    font-size: 15px;
    line-height: 1.45;
    color: #0c2349;
  }}
  .signature {{
    position: absolute;
    width: 420px;
    font-size: 15px;
    line-height: 1.4;
    color: #0c2349;
  }}
  .signature img {{
    width: 180px;
    height: auto;
    display: block;
  }}
  .signature-name {{
    margin-top: 10px;
    font-weight: 700;
  }}
  .signature-id {{
    color: #4b5563;
    font-size: 14px;
  }}
  .conditions h4 {{
    margin: 0 0 12px 0;
    font-size: 16px;
    font-weight: 800;
  }}
  .conditions ul {{
    margin: 0;
    padding-left: 18px;
    list-style: none;
  }}
  .conditions li::before {{
    content: "• ";
    color: #0c2349;
  }}

</style>
<div class="quote-page" id="quote-root">
  <div class="logo" style="left:{logo_left}px;top:{logo_top}px;width:{logo_box_width}px;height:{logo_box_height}px;">
    {"<img src='data:image/png;base64," + logo_b64 + "' alt='logo' style='width:" + str(logo_width) + "px;height:" + str(logo_height) + "px;' />" if logo_b64 else ""}
  </div>
  {header_repeats}
  <div class="header-info" style="left:{header_left}px;top:{header_top}px;width:{header_width}px;height:{header_height}px;">
    <div class="empresa">{html.escape(empresa)}</div>
    <div class="datos">{contacto_html}</div>
  </div>
  <div class="title" style="top:{title_top}px;left:{title_left}px;">Cotización</div>
  <div class="title-meta" style="top:{title_meta_top}px;left:{title_meta_left}px;">N.º cotización: <strong>{html.escape(numero)}</strong><br>Fecha: {fecha_cot.strftime('%Y-%m-%d')}</div>
  <div class="columns" style="top:{columns_top}px;left:{columns_left}px;">
    <div class="block">
      <h4>Datos del Cliente</h4>
      <div>{html.escape(cliente or '-')}</div>
      <div>{html.escape(direccion or '-')}</div>
      <div>{html.escape(cliente_ruc_text)}</div>
    </div>
    <div class="block">
      <h4>Datos del Emisor</h4>
      <div class="emisor-name">{html.escape(empresa)}</div>
      {"<div class=\"contacto\">" + contacto_html + "</div>" if contacto_html else ""}
    </div>
  </div>
  <div class="table-wrap" style="top:{table_top}px;left:{table_left}px;">
    <table class="items">
      <thead>
        <tr>
          <th>Producto</th>
          <th style="width:120px;">Cantidad</th>
          <th style="width:180px;">Precio unitario</th>
          <th style="width:180px;">Importe</th>
        </tr>
      </thead>
      <tbody>
        {sample_rows}
      </tbody>
    </table>
  </div>
  <div class="totals" style="top:{totals_top}px;right:{totals_right}px;">
    <div><span>Subtotal</span><span>{_format_money(subtotal)}</span></div>
    <div><span>Impuestos ({impuesto_pct:.2f}%)</span><span>{_format_money(impuesto)}</span></div>
    <div class="total"><span>TOTAL</span><span>{_format_money(total)}</span></div>
  </div>
  {extra_html}
  <div class="conditions" style="top:{conditions_top}px;left:{conditions_left}px;">
    <h4>CONDICIONES</h4>
    <ul>
      {condiciones_html}
    </ul>
  </div>
  {signature_html}
</div>
    """


def _build_budget_html(
    empresa: str,
    numero: str,
    fecha_cot: date,
    presupuesto_df: pd.DataFrame,
    costo_interno: float,
    factor_ganancia: float,
    precio_cotizar: float,
    ganancia: float,
    financiamiento_tipo: str,
    financiamiento_interes_pct: float,
    costo_financiamiento: float,
    ganancia_neta: float,
    tiempo_inversion: float,
    inversion_etapa_1: float,
    tiempo_inicio_ejecucion_presentacion: float,
    inversion_etapa_intermedia: float,
    tiempo_cobro: float,
    inversion_etapa_2: float,
) -> str:
    rows = []
    for _, row in presupuesto_df.iterrows():
        rows.append(
            f"""
            <tr>
              <td>{html.escape(str(row.get('producto_servicio', '') or ''))}</td>
              <td style="text-align:center;">{row.get('cantidad', 0):,.0f}</td>
              <td style="text-align:right;">{_format_money(row.get('precio_unitario', 0))}</td>
              <td style="text-align:right;">{_format_money(row.get('importe', 0))}</td>
            </tr>
            """
        )
    if not rows:
        rows.append(
            "<tr><td colspan=\"4\" style=\"text-align:center;color:#64748b;\">Sin items de presupuesto.</td></tr>"
        )
    tiempo_total = tiempo_inversion + tiempo_inicio_ejecucion_presentacion + tiempo_cobro
    tiempo_meses = tiempo_total / 30 if tiempo_total else 0.0

    return f"""
<!doctype html>
<html>
<head>
<meta charset="utf-8" />
<title>Presupuesto {numero}</title>
<style>
  body {{ font-family: Arial, sans-serif; margin: 32px; color: #0f172a; }}
  h1 {{ font-size: 22px; margin: 0 0 8px 0; }}
  h2 {{ font-size: 16px; margin: 24px 0 10px 0; }}
  .meta {{ color: #475569; font-size: 13px; margin-bottom: 16px; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 13px; }}
  th, td {{ border: 1px solid #cbd5f5; padding: 8px; vertical-align: top; }}
  th {{ background: #1c336a; color: #ffffff; text-align: left; }}
  .summary {{ margin-top: 18px; font-size: 14px; }}
  .summary div {{ display: flex; justify-content: space-between; margin-bottom: 6px; }}
  .summary .total {{ font-weight: 700; font-size: 15px; }}
</style>
</head>
<body>
  <h1>Presupuesto interno</h1>
  <div class="meta">Cotizacion: {html.escape(numero)} | Empresa: {html.escape(empresa)} | Fecha: {fecha_cot.strftime('%Y-%m-%d')}</div>
  <table>
    <thead>
      <tr>
        <th>Producto / Servicio</th>
        <th style="width:120px;">Cantidad</th>
        <th style="width:160px;">Precio unitario</th>
        <th style="width:160px;">Subtotal</th>
      </tr>
    </thead>
    <tbody>
      {''.join(rows)}
    </tbody>
  </table>
    <div class="summary">
    <div><span>Costo interno</span><span>{_format_money(costo_interno)}</span></div>
    <div><span>Factor de ganancia</span><span>{factor_ganancia:.2f}</span></div>
    <div><span>Precio a cotizar</span><span>{_format_money(precio_cotizar)}</span></div>
    <div><span>Ganancia</span><span>{_format_money(ganancia)}</span></div>
    <div><span>Financiamiento</span><span>{html.escape(financiamiento_tipo)} ({financiamiento_interes_pct:.2f}% mensual)</span></div>
    <div><span>Tiempo inversion→inicio ejecucion</span><span>{tiempo_inversion:.0f} dias</span></div>
    <div><span>Inversion etapa inversion→inicio ejecucion</span><span>{_format_money(inversion_etapa_1)}</span></div>
    <div><span>Tiempo inicio ejecucion→presentacion</span><span>{tiempo_inicio_ejecucion_presentacion:.0f} dias</span></div>
    <div><span>Inversion etapa inicio ejecucion→presentacion</span><span>{_format_money(inversion_etapa_intermedia)}</span></div>
    <div><span>Tiempo presentacion→cobro</span><span>{tiempo_cobro:.0f} dias</span></div>
    <div><span>Inversion etapa presentacion→cobro</span><span>{_format_money(inversion_etapa_2)}</span></div>
    <div><span>Costo financiamiento</span><span>{_format_money(costo_financiamiento)}</span></div>
    <div><span>Ganancia neta</span><span>{_format_money(ganancia_neta)}</span></div>
    <div><span>Tiempo recuperacion</span><span>{tiempo_total:.0f} dias (~{tiempo_meses:.1f} meses)</span></div>
  </div>
</body>
</html>
"""

def _render_pdf_component(
    html_body: str,
    filename: str,
    preview_scale: float = 0.75,
    pdf_max_pages: int = 2,
    render_scale: float = 1.5,
    jpeg_quality: float = 0.85,
) -> None:
    """Renderiza la vista previa y un botón JS para exportar a PDF usando html2canvas + jsPDF."""
    preview_height = 3200
    component_html = f"""
    <style>
      html, body {{
        margin: 0;
        padding: 0;
        background: #ffffff;
      }}
      .preview-shell {{
        width: 100%;
        display: flex;
        justify-content: center;
        overflow: auto;
        background: #ffffff;
      }}
      .preview-scale {{
        display: inline-block;
        overflow: hidden;
      }}
      .preview-scale .quote-page {{
        transform: scale({preview_scale});
        transform-origin: top left;
      }}
    </style>
    <div class="preview-shell">
      <div class="preview-scale">{html_body}</div>
    </div>
    <div id="pdf-clone-host" style="position: fixed; left: -100000px; top: 0;"></div>
    <div style="margin: 10px 0 16px 0;">
      <button id="btn-download" style="
        background: linear-gradient(135deg, #2563eb, #22c55e);
        color: white; border: none; padding: 10px 14px; border-radius: 10px;
        font-weight: 700; cursor: pointer; box-shadow: 0 8px 24px rgba(34,197,94,0.25);
      ">Descargar PDF</button>
    </div>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/html2canvas/1.4.1/html2canvas.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/jspdf/2.5.1/jspdf.umd.min.js"></script>
    <script>
      const previewScale = {preview_scale};
      const maxPages = {pdf_max_pages};
      const renderScale = {render_scale};
      const jpegQuality = {jpeg_quality};
      const previewWrapper = document.querySelector(".preview-scale");
      const previewQuote = document.querySelector(".preview-scale .quote-page");
      const syncPreviewSize = () => {{
        if (!previewWrapper || !previewQuote) return;
        previewWrapper.style.width = (previewQuote.offsetWidth * previewScale) + "px";
        previewWrapper.style.height = (previewQuote.offsetHeight * previewScale) + "px";
      }};
      window.requestAnimationFrame(syncPreviewSize);
      const btn = document.getElementById("btn-download");
      btn?.addEventListener("click", () => {{
        const root = document.getElementById("quote-root");
        const host = document.getElementById("pdf-clone-host");
        if (!root || !host) return;

        const clone = root.cloneNode(true);
        clone.removeAttribute("id");
        clone.style.transform = "none";
        clone.style.position = "relative";
        clone.style.left = "0";
        clone.style.top = "0";
        host.innerHTML = "";
        host.appendChild(clone);

        const render = () => {{
          html2canvas(clone, {{ scale: 2, useCORS: true, backgroundColor: "#ffffff" }}).then(canvas => {{
            const imgData = canvas.toDataURL("image/jpeg", jpegQuality);
            const pdf = new jspdf.jsPDF("p", "pt", "a4");
            const pageWidth = pdf.internal.pageSize.getWidth();
            const pageHeight = pdf.internal.pageSize.getHeight();
            const imgWidth = pageWidth;
            const imgHeight = canvas.height * (pageWidth / canvas.width);
            let heightLeft = imgHeight;
            let position = 0;

            pdf.addImage(imgData, "PNG", 0, position, imgWidth, imgHeight);
            heightLeft -= pageHeight;

            while (heightLeft > 0) {{
              position -= pageHeight;
              pdf.addPage();
              pdf.addImage(imgData, "PNG", 0, position, imgWidth, imgHeight);
              heightLeft -= pageHeight;
            }}
            pdf.save("{filename}");
            host.innerHTML = "";
          }});
        }};

        if (document.fonts && document.fonts.ready) {{
          document.fonts.ready.then(render);
        }} else {{
          render();
        }}
      }});
    </script>
    """
    components.html(component_html, height=preview_height, scrolling=True)


# ---- Configuración de empresas (membrete) ----
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS_DIR = os.path.join(os.path.dirname(BASE_DIR), "assets")

# Prefer paths proporcionados, luego assets de respaldo
RS_LOGO_PATH = os.path.join(ASSETS_DIR, "Logo RS Engineering.png")
RIR_LOGO_PATH = os.path.join(ASSETS_DIR, "Logo RIR Medical.png")
RS_LOGO_FALLBACK = os.path.join(ASSETS_DIR, "rs.png.png")
RIR_LOGO_FALLBACK = os.path.join(ASSETS_DIR, "rir.png.png")
BACKGROUND_PATH = os.path.join(ASSETS_DIR, "Fondo.png")
FIRMA_PATH = os.path.join(ASSETS_DIR, "firma.png")
BACKGROUND_B64 = _load_logo_b64(BACKGROUND_PATH)
FIRMA_B64 = _load_logo_b64(FIRMA_PATH)
COMPANIES = {
    "RS Engineering": {
        "color": "#0f172a",
        "accent": "#0e4aa0",
        "logo_b64": _load_logo_b64(RS_LOGO_PATH, RS_LOGO_FALLBACK),
        "background_b64": BACKGROUND_B64,
        "logo_box_width": 440,
        "logo_box_height": 440,
        "logo_width": 420,
        "logo_height": 420,
        "logo_left": 20,
        "logo_top": 80,
        "header_left": 430,
        "header_top": 80,
        "header_height": 440,
        "content_offset_y": 200,
        "contacto_html": """<div style='text-align:left; line-height:1.35;'>
        R.U.C. 9-740-624 / DV: 80<br>
        PH Bonanza Plaza, Bella Vista<br>
        TELÉFONO: +507 68475616<br>
        EMAIL: rodrigojesus-@hotmail.com
        </div>""",
    },
    "RIR Medical": {
        "color": "#1d4ed8",
        "accent": "#22c55e",
        "logo_b64": _load_logo_b64(RIR_LOGO_PATH, RIR_LOGO_FALLBACK),
        "background_b64": BACKGROUND_B64,
        "logo_box_width": 320,
        "logo_box_height": 170,
        "logo_width": 310,
        "logo_height": 166,
        "logo_left": 90,
        "logo_top": 100,
        "header_left": 430,
        "header_top": 100,
        "header_height": 170,
        "content_offset_y": 160,
        "contacto_html": """<div style='text-align:left; line-height:1.35;'>
        RUC: 155750585-2-2024 DV40<br>
        PH Bonanza Plaza, Bella Vista<br>
        TELÉFONO: +507 68475616<br>
        Email: info@rirmedical.com
        </div>""",
    },
}


# ---- UI principal ----
st.title("Generador de cotizaciones")

sheet_id = st.secrets.get("app", {}).get("SHEET_ID")
sheet_error = None
cot_df = pd.DataFrame(columns=COT_COLUMNS)
client = None
creds = None
if sheet_id:
    try:
        client, creds = get_client()
        if "cotizaciones_cache_token" not in st.session_state:
            st.session_state["cotizaciones_cache_token"] = uuid.uuid4().hex
        token = st.session_state["cotizaciones_cache_token"]
        cot_df = _normalize_cotizaciones_df(_load_cotizaciones_cached(sheet_id, token))
    except Exception as exc:
        sheet_error = str(exc)
else:
    sheet_error = "No hay SHEET_ID configurado en st.secrets['app']."

EDIT_KEY = "cotizacion_edit"
if EDIT_KEY not in st.session_state:
    st.session_state[EDIT_KEY] = None
PENDING_EDIT_KEY = "cotizacion_pending_edit_id"
PENDING_TAB_KEY = "cotizacion_pending_tab"
PENDING_DUPLICATE_KEY = "cotizacion_pending_duplicate_id"

items_state_key = "cotizacion_privada_items_data"
presupuesto_state_key = "cotizacion_presupuesto_items_data"


def _apply_edit_state(row: dict) -> None:
    st.session_state[EDIT_KEY] = row
    st.session_state["cot_empresa"] = row.get("empresa") or "RS Engineering"
    st.session_state["cot_cliente"] = row.get("cliente_nombre", "")
    st.session_state["cot_direccion"] = row.get("cliente_direccion", "")
    st.session_state["cot_cliente_ruc"] = row.get("cliente_ruc", "")
    st.session_state["cot_cliente_dv"] = row.get("cliente_dv", "")
    st.session_state["cot_detalles_extra"] = row.get("detalles_extra", "")
    st.session_state["cot_numero"] = row.get("numero_cotizacion", "")

    fecha_val = row.get("fecha_cotizacion") or ""
    fecha_dt = None
    if isinstance(fecha_val, str) and fecha_val:
        try:
            fecha_dt = datetime.fromisoformat(fecha_val).date()
        except ValueError:
            fecha_dt = None
    st.session_state["cot_fecha"] = fecha_dt or date.today()

    try:
        items = json.loads(row.get("items_json") or "[]")
        if not isinstance(items, list):
            items = []
    except Exception:
        items = []
    if not items:
        items = [{"producto_servicio": "Producto o servicio", "cantidad": 1, "precio_unitario": 100.0}]
    st.session_state[items_state_key] = items
    try:
        presupuesto_items = json.loads(row.get("presupuesto_items_json") or "[]")
        if not isinstance(presupuesto_items, list):
            presupuesto_items = []
    except Exception:
        presupuesto_items = []
    if not presupuesto_items:
        presupuesto_items = [
            {"producto_servicio": "Detalle", "cantidad": 1, "precio_unitario": 0.0}
        ]
    st.session_state[presupuesto_state_key] = presupuesto_items

    factor_val = row.get("presupuesto_factor_ganancia")
    try:
        factor_val = float(factor_val)
    except (TypeError, ValueError):
        factor_val = 1.3
    st.session_state["cot_presupuesto_factor"] = factor_val

    t_inv = row.get("presupuesto_t_inversion_presentacion")
    try:
        t_inv = float(t_inv)
    except (TypeError, ValueError):
        t_inv = 0.0
    st.session_state["cot_presupuesto_t_inversion"] = t_inv

    t_intermedio = row.get("presupuesto_t_inicio_ejecucion_presentacion")
    try:
        t_intermedio = float(t_intermedio)
    except (TypeError, ValueError):
        t_intermedio = 0.0
    st.session_state["cot_presupuesto_t_intermedio"] = t_intermedio

    t_cobro = row.get("presupuesto_t_presentacion_cobro")
    try:
        t_cobro = float(t_cobro)
    except (TypeError, ValueError):
        t_cobro = 0.0
    st.session_state["cot_presupuesto_t_cobro"] = t_cobro

    inv_etapa_1 = row.get("presupuesto_inversion_etapa_1")
    try:
        inv_etapa_1 = float(inv_etapa_1)
    except (TypeError, ValueError):
        inv_etapa_1 = 0.0
    st.session_state["cot_presupuesto_inv_etapa_1"] = inv_etapa_1

    inv_etapa_intermedia = row.get("presupuesto_inversion_etapa_intermedia")
    try:
        inv_etapa_intermedia = float(inv_etapa_intermedia)
    except (TypeError, ValueError):
        inv_etapa_intermedia = 0.0
    st.session_state["cot_presupuesto_inv_etapa_intermedia"] = inv_etapa_intermedia

    inv_etapa_2 = row.get("presupuesto_inversion_etapa_2")
    try:
        inv_etapa_2 = float(inv_etapa_2)
    except (TypeError, ValueError):
        inv_etapa_2 = 0.0
    st.session_state["cot_presupuesto_inv_etapa_2"] = inv_etapa_2

    fin_tipo = row.get("presupuesto_financiamiento_tipo") or "Dinero propio"
    if fin_tipo not in ("Dinero propio", "Prestamo"):
        fin_tipo = "Dinero propio"
    st.session_state["cot_presupuesto_fin_tipo"] = fin_tipo
    fin_interes = row.get("presupuesto_financiamiento_interes_pct")
    try:
        fin_interes = float(fin_interes)
    except (TypeError, ValueError):
        fin_interes = 2.5
    st.session_state["cot_presupuesto_fin_interes"] = fin_interes

    try:
        condiciones = json.loads(row.get("condiciones_json") or "{}")
    except Exception:
        condiciones = {}

    st.session_state["cot_vigencia"] = condiciones.get("Vigencia") or row.get("vigencia") or "15 días"
    st.session_state["cot_forma_pago"] = condiciones.get("Condicion de pago") or row.get("forma_pago") or "Credito"
    st.session_state["cot_entrega"] = condiciones.get("Entrega") or row.get("entrega") or "15 días hábiles"
    st.session_state["cot_lugar_entrega"] = (condiciones.get("Lugar de entrega") or row.get("lugar_entrega") or "")

    impuesto_val = row.get("impuesto_pct")
    try:
        impuesto_val = float(impuesto_val)
    except (TypeError, ValueError):
        impuesto_val = 7.0
    st.session_state["cot_impuesto"] = impuesto_val


def _clear_edit_state() -> None:
    st.session_state[EDIT_KEY] = None


def _apply_duplicate_state(row: dict, cotizaciones_df: pd.DataFrame) -> None:
    _apply_edit_state(row)
    _clear_edit_state()
    empresa_sel = st.session_state.get("cot_empresa") or row.get("empresa") or "RS Engineering"
    prefijo = COT_PREFIX.get(empresa_sel, "GEN")
    seq = _next_sequence(cotizaciones_df, prefijo)
    numero_auto = _build_numero_cot(prefijo, seq)
    st.session_state["cot_numero"] = numero_auto
    st.session_state["cot_numero_pref"] = prefijo


TAB_OPTIONS = [
    "Cotización - Panamá Compra",
    "Cotizacion - Estandar",
    "Historial de cotizaciones",
    LP_DOC_TAB_NAME,
]
pending_tab = st.session_state.pop(PENDING_TAB_KEY, None)
if pending_tab in TAB_OPTIONS:
    st.session_state["cotizaciones_tab"] = pending_tab
if st.session_state.get("cotizaciones_tab") == "Cotización - Privada":
    st.session_state["cotizaciones_tab"] = "Cotizacion - Estandar"
if st.session_state.get("cotizaciones_tab") not in TAB_OPTIONS:
    st.session_state["cotizaciones_tab"] = TAB_OPTIONS[0]

active_tab = st.segmented_control(
    "Secciones",
    TAB_OPTIONS,
    key="cotizaciones_tab",
    label_visibility="collapsed",
)

if active_tab == "Cotización - Panamá Compra":
    st.subheader("Generar cotización desde Panamá Compra")
    if "pc_presupuesto_items_data" not in st.session_state:
        st.session_state["pc_presupuesto_items_data"] = [
            {"producto_servicio": "Detalle", "cantidad": 1, "precio_unitario": 0.0},
        ]
    if "pc_presupuesto_factor" not in st.session_state:
        st.session_state["pc_presupuesto_factor"] = 1.3
    if "pc_presupuesto_t_inversion" not in st.session_state:
        st.session_state["pc_presupuesto_t_inversion"] = 0.0
    if "pc_presupuesto_t_intermedio" not in st.session_state:
        st.session_state["pc_presupuesto_t_intermedio"] = 0.0
    if "pc_presupuesto_t_cobro" not in st.session_state:
        st.session_state["pc_presupuesto_t_cobro"] = 0.0
    if "pc_presupuesto_inv_etapa_1" not in st.session_state:
        st.session_state["pc_presupuesto_inv_etapa_1"] = 0.0
    if "pc_presupuesto_inv_etapa_intermedia" not in st.session_state:
        st.session_state["pc_presupuesto_inv_etapa_intermedia"] = 0.0
    if "pc_presupuesto_inv_etapa_2" not in st.session_state:
        st.session_state["pc_presupuesto_inv_etapa_2"] = 0.0
    if "pc_presupuesto_fin_tipo" not in st.session_state:
        st.session_state["pc_presupuesto_fin_tipo"] = "Dinero propio"
    if "pc_presupuesto_fin_interes" not in st.session_state:
        st.session_state["pc_presupuesto_fin_interes"] = 2.5
    if "pc_cot_empresa" not in st.session_state:
        st.session_state["pc_cot_empresa"] = "RS"
    if "pc_cot_itbms" not in st.session_state:
        st.session_state["pc_cot_itbms"] = True
    if "pc_cot_precio_auto" not in st.session_state:
        st.session_state["pc_cot_precio_auto"] = True
    if "pc_cot_precio_manual" not in st.session_state:
        st.session_state["pc_cot_precio_manual"] = not bool(st.session_state.get("pc_cot_precio_auto", True))
    if "pc_cot_precio" not in st.session_state:
        st.session_state["pc_cot_precio"] = 0.0
    if "pc_cot_cliente_ruc" not in st.session_state:
        st.session_state["pc_cot_cliente_ruc"] = ""
    if "pc_cot_cliente_dv" not in st.session_state:
        st.session_state["pc_cot_cliente_dv"] = ""

    st.markdown("### Presupuesto interno (base de participación)")
    presupuesto_display_df = _build_items_dataframe(
        pd.DataFrame(st.session_state["pc_presupuesto_items_data"])
    )
    presupuesto_raw = st.data_editor(
        presupuesto_display_df,
        num_rows="dynamic",
        use_container_width=True,
        key="pc_presupuesto_items_editor",
        column_config={
            "producto_servicio": st.column_config.TextColumn("Producto / Servicio", width="large", required=True),
            "cantidad": st.column_config.NumberColumn("Cantidad", min_value=0.0, step=1.0, required=True),
            "precio_unitario": st.column_config.NumberColumn(
                "Precio unitario", min_value=0.0, step=10.0, format="$%0.2f", required=True
            ),
            "importe": st.column_config.NumberColumn("Subtotal", format="$%0.2f", disabled=True),
        },
        disabled=["importe"],
        hide_index=True,
    )
    presupuesto_df_pc = _build_items_dataframe(pd.DataFrame(presupuesto_raw))
    st.session_state["pc_presupuesto_items_data"] = presupuesto_df_pc[
        ["producto_servicio", "cantidad", "precio_unitario"]
    ].to_dict(orient="records")
    costo_interno_pc = float(presupuesto_df_pc["importe"].sum())

    col_pb1, col_pb2, col_pb3, col_pb4 = st.columns([1, 1, 1, 1])
    with col_pb1:
        factor_ganancia_pc = st.number_input(
            "Factor de ganancia",
            min_value=0.0,
            step=0.05,
            key="pc_presupuesto_factor",
        )
    with col_pb2:
        tiempo_inversion_pc = st.number_input(
            "Tiempo inversion→inicio de ejecucion (dias)",
            min_value=0.0,
            step=1.0,
            key="pc_presupuesto_t_inversion",
        )
        inversion_etapa_1_pc = st.number_input(
            "Inversion requerida etapa 1",
            min_value=0.0,
            step=100.0,
            key="pc_presupuesto_inv_etapa_1",
        )
    with col_pb3:
        tiempo_intermedio_pc = st.number_input(
            "Tiempo inicio de ejecucion→presentacion (dias)",
            min_value=0.0,
            step=1.0,
            key="pc_presupuesto_t_intermedio",
        )
        inversion_etapa_intermedia_pc = st.number_input(
            "Inversion requerida etapa 2",
            min_value=0.0,
            step=100.0,
            key="pc_presupuesto_inv_etapa_intermedia",
        )
    with col_pb4:
        tiempo_cobro_pc = st.number_input(
            "Tiempo presentacion→cobro (dias)",
            min_value=0.0,
            step=1.0,
            key="pc_presupuesto_t_cobro",
        )
        inversion_etapa_2_pc = st.number_input(
            "Inversion requerida etapa 3",
            min_value=0.0,
            step=100.0,
            key="pc_presupuesto_inv_etapa_2",
        )

    col_pf1, col_pf2 = st.columns([1, 1])
    with col_pf1:
        financiamiento_tipo_pc = st.selectbox(
            "Financiamiento",
            ["Dinero propio", "Prestamo"],
            key="pc_presupuesto_fin_tipo",
        )
    with col_pf2:
        financiamiento_interes_pct_pc = st.number_input(
            "Interes mensual (%)",
            min_value=0.0,
            step=0.1,
            key="pc_presupuesto_fin_interes",
        )

    tiempo_recuperacion_pc = tiempo_inversion_pc + tiempo_intermedio_pc + tiempo_cobro_pc
    tiempo_recuperacion_meses_pc = tiempo_recuperacion_pc / 30 if tiempo_recuperacion_pc else 0.0
    costo_financiamiento_pc = 0.0
    if financiamiento_tipo_pc == "Prestamo":
        tasa_mensual_pc = financiamiento_interes_pct_pc / 100.0
        meses_1 = tiempo_inversion_pc / 30 if tiempo_inversion_pc else 0.0
        meses_2 = tiempo_intermedio_pc / 30 if tiempo_intermedio_pc else 0.0
        meses_3 = tiempo_cobro_pc / 30 if tiempo_cobro_pc else 0.0
        if (
            inversion_etapa_1_pc <= 0
            and inversion_etapa_intermedia_pc <= 0
            and inversion_etapa_2_pc <= 0
        ):
            costo_financiamiento_pc = costo_interno_pc * tasa_mensual_pc * tiempo_recuperacion_meses_pc
        else:
            costo_financiamiento_pc = (
                inversion_etapa_1_pc * tasa_mensual_pc * meses_1
                + inversion_etapa_intermedia_pc * tasa_mensual_pc * meses_2
                + inversion_etapa_2_pc * tasa_mensual_pc * meses_3
            )
    costo_total_final_pc = costo_interno_pc + costo_financiamiento_pc
    precio_cotizar_pc = costo_total_final_pc * factor_ganancia_pc
    ganancia_pc = precio_cotizar_pc - costo_total_final_pc
    ganancia_neta_pc = precio_cotizar_pc - costo_total_final_pc

    st.markdown(
        f"**Resumen presupuesto:** Costo interno {_format_money(costo_interno_pc)} | "
        f"Base final (interno + financiamiento) {_format_money(costo_total_final_pc)} | "
        f"Precio sugerido participación {_format_money(precio_cotizar_pc)} | "
        f"Ganancia {_format_money(ganancia_pc)} | "
        f"Costo financiamiento {_format_money(costo_financiamiento_pc)} | "
        f"Ganancia neta {_format_money(ganancia_neta_pc)} | "
        f"Tiempo recuperacion {tiempo_recuperacion_pc:.0f} dias (~{tiempo_recuperacion_meses_pc:.1f} meses)"
    )

    st.markdown("### Parámetros de generación")
    col_a, col_b, col_c = st.columns([2.1, 1, 1])
    with col_a:
        enlace_pc = st.text_input("Enlace de Panamá Compra", key="pc_cot_enlace")
    with col_b:
        empresa_sel = st.selectbox("Empresa", ["RS", "RIR"], key="pc_cot_empresa")
    with col_c:
        paga_itbms = st.checkbox("Aplica ITBMS (7%)", key="pc_cot_itbms")

    col_ruc, col_dv = st.columns([1.4, 1])
    with col_ruc:
        cliente_ruc_pc = st.text_input("RUC cliente (opcional)", key="pc_cot_cliente_ruc")
    with col_dv:
        cliente_dv_pc = st.text_input("DV (opcional)", key="pc_cot_cliente_dv")

    manual_price_mode = st.toggle(
        "Precio de participación manual",
        key="pc_cot_precio_manual",
        help="Si está apagado, se usa automáticamente el precio sugerido por presupuesto.",
    )
    if not manual_price_mode:
        st.session_state["pc_cot_precio"] = round(float(precio_cotizar_pc), 2)
    precio_part = st.number_input(
        "Precio de participación manual",
        min_value=0.0,
        step=10.0,
        format="%0.2f",
        key="pc_cot_precio",
        disabled=not bool(manual_price_mode),
    )
    if not manual_price_mode:
        st.caption(f"Precio aplicado (presupuesto): {_format_money(precio_cotizar_pc)}")

    if st.button("Generar cotización (Panamá Compra)"):
        if not enlace_pc.strip():
            st.warning("Debes pegar el enlace de Panamá Compra.")
        elif float(precio_part) <= 0:
            st.warning("El precio de participación debe ser mayor a 0.")
        else:
            try:
                client_manual, _ = get_client()
                _ensure_pc_config_job(client_manual)
                payload = {
                    "enlace": enlace_pc.strip(),
                    "precio_participacion": float(precio_part),
                    "paga_itbms": bool(paga_itbms),
                    "empresa": empresa_sel.strip().lower(),
                    "cliente_ruc": str(cliente_ruc_pc or "").strip(),
                    "cliente_dv": str(cliente_dv_pc or "").strip(),
                }
                request_id = _append_manual_request(client_manual, payload)
                st.session_state["pc_cot_request_id"] = request_id
                st.session_state["pc_cot_payload"] = payload
                st.session_state.pop("pc_cot_processed_request_id", None)
                st.session_state.pop("pc_cot_processed_file_id", None)
                st.session_state.pop("pc_cot_final_excel_bytes", None)
                st.session_state.pop("pc_cot_final_excel_name", None)
                st.success("Solicitud enviada. El scraper iniciará el proceso.")
            except Exception as exc:
                st.error(f"No se pudo enviar la solicitud: {exc}")

    request_id = st.session_state.get("pc_cot_request_id")
    if request_id:
        try:
            client_manual, creds_manual = get_client()
            row = _fetch_manual_request(client_manual, request_id)
        except Exception as exc:
            row = None
            st.error(f"No se pudo consultar el estado: {exc}")

        if row:
            status = (row.get("status") or "").strip().lower()
            notes = (row.get("notes") or "").strip()
            progress_value = {
                "pending": 0.15,
                "enqueued": 0.35,
                "running": 0.75,
                "done": 1.0,
                "error": 1.0,
            }.get(status, 0.1)
            progress_text = {
                "pending": "Generando: solicitud recibida",
                "enqueued": "Generando: en cola de ejecución",
                "running": "Generando: scraping en ejecución",
                "done": "Generado: cotización lista",
                "error": "Error en la generación",
            }.get(status, "Generando...")
            st.progress(progress_value, text=progress_text)
            st.info(f"Estado actual: {status or 'pendiente'}")
            if notes:
                notes_display = (
                    notes.replace("Orquestador", "Scraping")
                    .replace("orquestador", "scraping")
                    .replace("Orchestrator", "Scraping")
                    .replace("orchestrator", "scraping")
                )
                st.caption(notes_display)
            if row.get("result_error"):
                st.error(row["result_error"])

            file_id = (row.get("result_file_id") or "").strip()
            if status == "done" and not file_id:
                st.warning("La solicitud terminó, pero aún no aparece el archivo. Actualiza el estado.")

            payload_row = _parse_manual_payload(row.get("payload", ""))
            empresa_from_payload = (
                str(payload_row.get("empresa") or st.session_state.get("pc_cot_empresa") or "RS").upper()
            )
            if empresa_from_payload not in {"RS", "RIR"}:
                empresa_from_payload = "RS"
            enlace_from_payload = str(payload_row.get("enlace") or st.session_state.get("pc_cot_enlace") or "")
            paga_itbms_payload = bool(payload_row.get("paga_itbms", st.session_state.get("pc_cot_itbms", True)))
            cliente_ruc_payload = str(
                payload_row.get("cliente_ruc", st.session_state.get("pc_cot_cliente_ruc", "")) or ""
            ).strip()
            cliente_dv_payload = str(
                payload_row.get("cliente_dv", st.session_state.get("pc_cot_cliente_dv", "")) or ""
            ).strip()
            try:
                precio_participacion_payload = float(
                    payload_row.get("precio_participacion", st.session_state.get("pc_cot_precio", 0.0))
                )
            except (TypeError, ValueError):
                precio_participacion_payload = float(st.session_state.get("pc_cot_precio", 0.0) or 0.0)
            costo_total_final_participacion = costo_interno_pc + costo_financiamiento_pc
            ganancia_participacion = precio_participacion_payload - costo_total_final_participacion
            ganancia_neta_participacion = precio_participacion_payload - costo_total_final_participacion

            processed_req = st.session_state.get("pc_cot_processed_request_id")
            processed_file = st.session_state.get("pc_cot_processed_file_id")
            if status == "done" and file_id and (processed_req != request_id or processed_file != file_id):
                try:
                    drive = _get_drive_client(creds_manual)
                    source_bytes = _download_drive_file(drive, file_id)
                    items_panama_df, titulo_excel, _, panama_meta = _extract_excel_items(source_bytes)
                    save_result = _save_panama_quote_to_history(
                        client=client_manual,
                        creds=creds_manual,
                        sheet_id=sheet_id,
                        cot_df=cot_df,
                        manual_request_id=request_id,
                        empresa_short=empresa_from_payload,
                        enlace_pc=enlace_from_payload,
                        titulo_excel=titulo_excel,
                        panama_meta=panama_meta,
                        items_panama_df=items_panama_df,
                        paga_itbms=paga_itbms_payload,
                        presupuesto_df=presupuesto_df_pc,
                        costo_interno=costo_interno_pc,
                        factor_ganancia=factor_ganancia_pc,
                        precio_cotizar=precio_participacion_payload,
                        ganancia=ganancia_participacion,
                        financiamiento_tipo=financiamiento_tipo_pc,
                        financiamiento_interes_pct=financiamiento_interes_pct_pc,
                        costo_financiamiento=costo_financiamiento_pc,
                        ganancia_neta=ganancia_neta_participacion,
                        tiempo_inversion=tiempo_inversion_pc,
                        inversion_etapa_1=inversion_etapa_1_pc,
                        tiempo_intermedio=tiempo_intermedio_pc,
                        inversion_etapa_intermedia=inversion_etapa_intermedia_pc,
                        tiempo_cobro=tiempo_cobro_pc,
                        inversion_etapa_2=inversion_etapa_2_pc,
                        manual_cliente_ruc=cliente_ruc_payload,
                        manual_cliente_dv=cliente_dv_payload,
                    )
                    st.session_state["pc_cot_final_excel_bytes"] = save_result["excel_bytes"]
                    st.session_state["pc_cot_final_excel_name"] = save_result["excel_name"]
                    st.session_state["pc_cot_final_numero"] = save_result["numero_cotizacion"]
                    st.session_state["pc_cot_processed_request_id"] = request_id
                    st.session_state["pc_cot_processed_file_id"] = file_id
                    st.session_state["cotizaciones_cache_token"] = uuid.uuid4().hex
                    st.success(
                        f"Cotización convertida a formato estándar y guardada en historial: {save_result['numero_cotizacion']}"
                    )
                except Exception as exc:
                    st.error(f"No se pudo procesar la cotización final: {exc}")

            if "pc_cot_auto_refresh" not in st.session_state:
                st.session_state["pc_cot_auto_refresh"] = status in {"pending", "enqueued", "running"}
            auto_refresh = st.checkbox(
                "Actualizar automáticamente (cada 5s)",
                key="pc_cot_auto_refresh",
            )
            if auto_refresh and status in {"pending", "enqueued", "running"}:
                time.sleep(5)
                st.rerun()

    if st.session_state.get("pc_cot_final_excel_bytes"):
        st.markdown("### Vista previa (Excel final)")
        preview_pc = _extract_standard_excel_preview(st.session_state["pc_cot_final_excel_bytes"])
        c1, c2, c3 = st.columns(3)
        with c1:
            st.caption("NÚMERO")
            st.write(preview_pc.get("numero") or st.session_state.get("pc_cot_final_numero") or "-")
            st.caption("CLIENTE")
            st.write(preview_pc.get("cliente") or "-")
        with c2:
            st.caption("FECHA")
            st.write(preview_pc.get("fecha") or "-")
            st.caption("RUC/DV")
            st.write(preview_pc.get("ruc_dv") or "-")
        with c3:
            st.caption("TÍTULO")
            st.write(preview_pc.get("titulo") or "-")

        preview_items_df = pd.DataFrame(preview_pc.get("items") or [])
        if not preview_items_df.empty:
            st.dataframe(
                preview_items_df,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Costo Unitario": st.column_config.NumberColumn("Costo Unitario", format="$%0.2f"),
                    "Total": st.column_config.NumberColumn("Total", format="$%0.2f"),
                },
            )
        st.markdown(
            f"**Totales en Excel:** Subtotal {_format_money(preview_pc.get('subtotal', 0.0))} | "
            f"Impuesto {_format_money(preview_pc.get('impuesto', 0.0))} | "
            f"Total {_format_money(preview_pc.get('total', 0.0))}"
        )
        st.download_button(
            "Descargar Excel de cotización",
            data=st.session_state["pc_cot_final_excel_bytes"],
            file_name=st.session_state.get("pc_cot_final_excel_name") or "cotizacion_panama_estandar.xlsx",
            mime=_guess_mime_from_filename(
                st.session_state.get("pc_cot_final_excel_name") or "cotizacion_panama_estandar.xlsx"
            ),
        )

if active_tab == LP_DOC_TAB_NAME:
    st.subheader("LP_ Doc_Generator")
    st.caption("Basado en plantillas de doc_gen. Al generar, en Drive se reemplazan archivos previos por nombre.")

    if "lp_doc_fecha" not in st.session_state:
        st.session_state["lp_doc_fecha"] = date.today()
    if "lp_doc_empresa" not in st.session_state:
        st.session_state["lp_doc_empresa"] = "RS"
    if "lp_doc_enlace" not in st.session_state:
        st.session_state["lp_doc_enlace"] = ""
    if "lp_doc_docs" not in st.session_state:
        st.session_state["lp_doc_docs"] = []
    if "lp_doc_links" not in st.session_state:
        st.session_state["lp_doc_links"] = {}
    if "lp_doc_zip_name" not in st.session_state:
        st.session_state["lp_doc_zip_name"] = ""
    if "lp_doc_zip_bytes" not in st.session_state:
        st.session_state["lp_doc_zip_bytes"] = b""
    if "lp_doc_folder_url" not in st.session_state:
        st.session_state["lp_doc_folder_url"] = ""

    col_lp1, col_lp2 = st.columns([2.6, 1])
    with col_lp1:
        enlace_lp = st.text_input("Enlace Panamá Compra", key="lp_doc_enlace")
    with col_lp2:
        empresa_lp = st.selectbox("Empresa", ["RS", "RIR"], key="lp_doc_empresa")
    st.caption(
        "Al generar, el sistema ejecuta scraping automáticamente del enlace (versión 3), "
        "extrae datos del acto y crea los documentos LP."
    )

    col_lp4, col_lp5 = st.columns([1, 1])
    with col_lp4:
        st.text_input(
            "Representante legal (Pacto)",
            key="lp_doc_representante_pacto",
            placeholder="Nombre para Pacto de integridad",
        )
        st.text_input(
            "Representante legal (Documentos)",
            key="lp_doc_representante_docs",
            placeholder="Nombre para resto de documentos",
        )
    with col_lp5:
        st.text_input("Cédula representante", key="lp_doc_cedula", placeholder="8-888-888")
        st.date_input("Fecha del documento", key="lp_doc_fecha")

    if st.button("Generar documentos LP"):
        representante_pacto = str(st.session_state.get("lp_doc_representante_pacto") or "").strip()
        representante_docs = str(st.session_state.get("lp_doc_representante_docs") or "").strip()
        cedula_lp = str(st.session_state.get("lp_doc_cedula") or "").strip()
        enlace_lp = str(st.session_state.get("lp_doc_enlace") or "").strip()

        missing = []
        if not enlace_lp:
            missing.append("Enlace Panamá Compra")
        if not representante_pacto:
            missing.append("Representante legal (Pacto)")
        if not representante_docs:
            missing.append("Representante legal (Documentos)")
        if not cedula_lp:
            missing.append("Cédula representante")

        if missing:
            st.warning("Completa estos campos: " + ", ".join(missing))
        else:
            try:
                progress_ph = st.empty()
                notes_ph = st.empty()
                progress_ph.progress(0.05, text="Generando: iniciando scraping de datos del acto")

                client_manual_lp, creds_manual_lp = get_client()
                _ensure_pc_config_job(client_manual_lp)
                payload_lp = {
                    "enlace": enlace_lp,
                    "precio_participacion": 1.0,
                    "paga_itbms": False,
                    "empresa": str(st.session_state.get("lp_doc_empresa") or "RS").strip().lower(),
                    "modo": "lp_doc_autofill",
                }
                lp_request_id = _append_manual_request(client_manual_lp, payload_lp)

                start_wait = time.time()
                timeout_seconds = 300
                lp_row = None
                lp_status = "pending"
                while time.time() - start_wait < timeout_seconds:
                    lp_row = _fetch_manual_request(client_manual_lp, lp_request_id)
                    lp_status = str((lp_row or {}).get("status") or "").strip().lower()
                    lp_notes = str((lp_row or {}).get("notes") or "").strip()
                    progress_val = {
                        "pending": 0.2,
                        "enqueued": 0.35,
                        "running": 0.7,
                        "done": 1.0,
                        "error": 1.0,
                    }.get(lp_status, 0.15)
                    progress_txt = {
                        "pending": "Generando: solicitud enviada",
                        "enqueued": "Generando: scraping en cola",
                        "running": "Generando: scraping en ejecución",
                        "done": "Generando: scraping finalizado",
                        "error": "Error en scraping",
                    }.get(lp_status, "Generando...")
                    progress_ph.progress(progress_val, text=progress_txt)
                    if lp_notes:
                        notes_ph.caption(
                            lp_notes.replace("Orquestador", "Scraping")
                            .replace("orquestador", "scraping")
                            .replace("Orchestrator", "Scraping")
                            .replace("orchestrator", "scraping")
                        )
                    if lp_status in {"done", "error"}:
                        break
                    time.sleep(5)

                if lp_status not in {"done"}:
                    if lp_status == "error":
                        error_msg = str((lp_row or {}).get("result_error") or "").strip()
                        raise RuntimeError(error_msg or "Falló el scraping para extraer datos del acto.")
                    raise RuntimeError("El scraping tardó demasiado en completar (timeout de 5 minutos).")

                lp_file_id = str((lp_row or {}).get("result_file_id") or "").strip()
                if not lp_file_id:
                    raise RuntimeError("El scraping terminó, pero no devolvió archivo para extraer datos.")

                lp_drive = _get_drive_client(creds_manual_lp)
                lp_source_bytes = _download_drive_file(lp_drive, lp_file_id)
                _, lp_titulo_excel, _, lp_meta = _extract_excel_items(lp_source_bytes)

                entidad_lp = str(lp_meta.get("entidad") or "").strip()
                titulo_lp = str(lp_titulo_excel or "").strip()
                numero_lp = (
                    str(lp_meta.get("numero_acto") or "").strip()
                    or _extract_numero_acto_from_link(enlace_lp)
                )
                lugar_lp = str(lp_meta.get("lugar_entrega") or "").strip()
                tiempo_lp = str(lp_meta.get("tiempo_entrega") or "").strip()

                missing_extracted = []
                if not entidad_lp:
                    missing_extracted.append("Entidad")
                if not titulo_lp:
                    missing_extracted.append("Título del acto")
                if not numero_lp:
                    missing_extracted.append("Número de acto")
                if not lugar_lp:
                    missing_extracted.append("Lugar de entrega")
                if not tiempo_lp:
                    missing_extracted.append("Tiempo de entrega")
                if missing_extracted:
                    raise RuntimeError(
                        "El scraping no pudo extraer: " + ", ".join(missing_extracted)
                    )

                empresa_full_lp = _company_full_from_short(st.session_state.get("lp_doc_empresa") or "RS")
                docs_generated = _build_lp_documents(
                    empresa_full=empresa_full_lp,
                    fecha_base=st.session_state.get("lp_doc_fecha") or date.today(),
                    representante_legal_pacto=representante_pacto,
                    representante_legal_documentos=representante_docs,
                    cedula=cedula_lp,
                    entidad=entidad_lp,
                    titulo=titulo_lp,
                    numero_acto=numero_lp,
                    lugar_entrega=lugar_lp,
                    tiempo_entrega=tiempo_lp,
                )

                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
                    for doc_info in docs_generated:
                        zf.writestr(doc_info["file_name"], doc_info["bytes"])

                st.session_state["lp_doc_docs"] = docs_generated
                st.session_state["lp_doc_zip_name"] = (
                    f"LP_Doc_Generator_{(st.session_state.get('lp_doc_empresa') or 'RS').upper()}_{numero_lp}.zip"
                )
                st.session_state["lp_doc_zip_bytes"] = zip_buffer.getvalue()

                links_map: dict[str, str] = {}
                folder_url = ""
                try:
                    if creds is None:
                        client, creds = get_client()
                    drive = _get_drive_client(creds)
                    lp_folder_id = _get_lp_doc_folder(drive, empresa_full_lp)
                    if lp_folder_id:
                        folder_url = f"https://drive.google.com/drive/folders/{lp_folder_id}"
                    for doc_info in docs_generated:
                        file_name = str(doc_info["file_name"])
                        existing_id = _find_file_in_folder(
                            drive,
                            folder_id=lp_folder_id,
                            filename=file_name,
                        )
                        upload = _upload_drive_binary(
                            drive,
                            lp_folder_id,
                            file_name,
                            doc_info["bytes"],
                            _guess_mime_from_filename(file_name),
                            existing_file_id=existing_id or None,
                        )
                        file_id = upload.get("id")
                        if file_id:
                            links_map[file_name] = f"https://drive.google.com/file/d/{file_id}/view"
                except Exception as drive_exc:
                    st.warning(f"Documentos generados localmente. No se pudieron sincronizar en Drive: {drive_exc}")

                st.session_state["lp_doc_links"] = links_map
                st.session_state["lp_doc_folder_url"] = folder_url
                st.success("Documentos LP generados correctamente.")
            except Exception as exc:
                st.error(f"No se pudieron generar los documentos LP: {exc}")

    docs_generated = st.session_state.get("lp_doc_docs") or []
    if docs_generated:
        st.markdown("### Documentos generados")
        folder_url = str(st.session_state.get("lp_doc_folder_url") or "").strip()
        if folder_url:
            st.markdown(f"Carpeta en Drive: [Abrir carpeta]({folder_url})")
        st.download_button(
            "Descargar paquete ZIP",
            data=st.session_state.get("lp_doc_zip_bytes") or b"",
            file_name=st.session_state.get("lp_doc_zip_name") or "LP_Doc_Generator.zip",
            mime="application/zip",
            key="lp_doc_zip_download",
        )
        link_map = st.session_state.get("lp_doc_links") or {}
        for idx, doc_info in enumerate(docs_generated, start=1):
            file_name = str(doc_info.get("file_name") or f"documento_{idx}.docx")
            cols = st.columns([1.8, 1.2, 1.2])
            with cols[0]:
                st.write(file_name)
            with cols[1]:
                st.download_button(
                    "Descargar .docx",
                    data=doc_info.get("bytes") or b"",
                    file_name=file_name,
                    mime=_guess_mime_from_filename(file_name),
                    key=f"lp_doc_download_{idx}_{file_name}",
                )
            with cols[2]:
                drive_url = link_map.get(file_name)
                if drive_url:
                    st.markdown(f"[Abrir en Drive]({drive_url})")
                else:
                    st.caption("Sin enlace de Drive")

if active_tab == "Cotizacion - Estandar":
    if sheet_error:
        st.warning(sheet_error)

    pending_dup = st.session_state.pop(PENDING_DUPLICATE_KEY, None)
    pending_id = st.session_state.pop(PENDING_EDIT_KEY, None)
    if pending_dup and not cot_df.empty:
        row_match = cot_df[cot_df["id"] == pending_dup]
        if not row_match.empty:
            _apply_duplicate_state(row_match.iloc[0].to_dict(), cot_df)
    elif pending_id and not cot_df.empty:
        row_match = cot_df[cot_df["id"] == pending_id]
        if not row_match.empty:
            _apply_edit_state(row_match.iloc[0].to_dict())

    edit_row = st.session_state.get(EDIT_KEY)
    if edit_row:
        st.info(f"Editando: {edit_row.get('numero_cotizacion', '')}")
        if st.button("Cancelar edición"):
            _clear_edit_state()
            st.rerun()

    if "cot_fecha" not in st.session_state:
        st.session_state["cot_fecha"] = date.today()
    if "cot_impuesto" not in st.session_state:
        st.session_state["cot_impuesto"] = 7.0
    if "cot_detalles_extra" not in st.session_state:
        st.session_state["cot_detalles_extra"] = ""
    if "cot_lugar_entrega" not in st.session_state:
        st.session_state["cot_lugar_entrega"] = ""
    if "cot_cliente_ruc" not in st.session_state:
        st.session_state["cot_cliente_ruc"] = ""
    if "cot_cliente_dv" not in st.session_state:
        st.session_state["cot_cliente_dv"] = ""
    if "cot_presupuesto_factor" not in st.session_state:
        st.session_state["cot_presupuesto_factor"] = 1.3
    if "cot_presupuesto_t_inversion" not in st.session_state:
        st.session_state["cot_presupuesto_t_inversion"] = 0.0
    if "cot_presupuesto_t_intermedio" not in st.session_state:
        st.session_state["cot_presupuesto_t_intermedio"] = 0.0
    if "cot_presupuesto_t_cobro" not in st.session_state:
        st.session_state["cot_presupuesto_t_cobro"] = 0.0
    if "cot_presupuesto_inv_etapa_1" not in st.session_state:
        st.session_state["cot_presupuesto_inv_etapa_1"] = 0.0
    if "cot_presupuesto_inv_etapa_intermedia" not in st.session_state:
        st.session_state["cot_presupuesto_inv_etapa_intermedia"] = 0.0
    if "cot_presupuesto_inv_etapa_2" not in st.session_state:
        st.session_state["cot_presupuesto_inv_etapa_2"] = 0.0
    if "cot_presupuesto_fin_tipo" not in st.session_state:
        st.session_state["cot_presupuesto_fin_tipo"] = "Dinero propio"
    if "cot_presupuesto_fin_interes" not in st.session_state:
        st.session_state["cot_presupuesto_fin_interes"] = 2.5

    st.subheader("Datos de la cotización")
    col_a, col_b, col_c = st.columns([1.2, 1, 1])
    with col_a:
        empresa = st.selectbox("Empresa", list(COMPANIES.keys()), key="cot_empresa")
        if sheet_error or not sheet_id or client is None:
            st.caption("Catálogo de clientes no disponible.")
        else:
            cliente_id, cliente_nombre = client_selector(client, sheet_id, key="cot_catalogo")
            if cliente_nombre:
                st.session_state["cot_cliente"] = cliente_nombre
                st.session_state["cot_cliente_id"] = cliente_id
        cliente = st.text_input("Nombre del cliente", key="cot_cliente")
        direccion = st.text_area("Dirección del cliente", height=70, key="cot_direccion")
        col_ruc, col_dv = st.columns([2, 1])
        with col_ruc:
            cliente_ruc = st.text_input("RUC del cliente", key="cot_cliente_ruc")
        with col_dv:
            cliente_dv = st.text_input("DV", key="cot_cliente_dv")

        if not sheet_error and sheet_id and client is not None:
            with st.expander("Cliente no registrado? Agregar al catalogo", expanded=False):
                default_emp = CLIENT_EMPRESA_MAP.get(empresa, CLIENT_EMPRESA_OPTIONS[0])
                try:
                    default_idx = CLIENT_EMPRESA_OPTIONS.index(default_emp)
                except ValueError:
                    default_idx = 0
                nuevo_nombre = st.text_input(
                    "Nombre del nuevo cliente",
                    value=cliente or "",
                    key="cot_cliente_nuevo",
                )
                nuevo_emp = st.selectbox(
                    "Empresa (cliente)",
                    CLIENT_EMPRESA_OPTIONS,
                    index=default_idx,
                    key="cot_cliente_empresa",
                )
                if st.button("Guardar cliente", key="cot_cliente_guardar"):
                    if not nuevo_nombre.strip():
                        st.warning("Debes indicar el nombre del cliente.")
                    else:
                        try:
                            nuevo_id, created = _create_cliente_in_sheet(
                                client,
                                sheet_id,
                                nuevo_nombre,
                                nuevo_emp,
                            )
                            st.session_state["cot_cliente"] = nuevo_nombre.strip()
                            st.session_state["cot_cliente_id"] = nuevo_id
                            if created:
                                st.toast(f"Cliente creado: {nuevo_id}")
                            else:
                                st.toast("El cliente ya existia en el catalogo.")
                            st.rerun()
                        except Exception as exc:
                            st.error(f"No se pudo crear el cliente: {exc}")
    with col_b:
        prefijo = COT_PREFIX.get(empresa, "GEN")
        seq = _next_sequence(cot_df, prefijo)
        numero_auto = _build_numero_cot(prefijo, seq)
        if edit_row:
            numero_auto = edit_row.get("numero_cotizacion") or numero_auto
        if not edit_row:
            if st.session_state.get("cot_numero_pref") != prefijo:
                st.session_state["cot_numero"] = numero_auto
                st.session_state["cot_numero_pref"] = prefijo
        numero_cot = st.text_input("Número de cotización", key="cot_numero", disabled=True)
        fecha_cot = st.date_input("Fecha", key="cot_fecha")
        impuesto_pct = st.number_input("Impuesto (%)", min_value=0.0, max_value=25.0, step=0.5, key="cot_impuesto")
    with col_c:
        vigencia = st.text_input("Vigencia de la oferta", value="15 días", key="cot_vigencia")
        forma_pago = st.selectbox("Condicion de pago", ["Credito", "Contado"], index=0, key="cot_forma_pago")
        entrega = st.text_input("Entrega", value="15 días hábiles", key="cot_entrega")
        lugar_entrega = st.text_input("Lugar de entrega", key="cot_lugar_entrega")

    st.markdown("### Ítems de la cotización")
    if items_state_key not in st.session_state:
        st.session_state[items_state_key] = [
            {"producto_servicio": "Producto o servicio", "cantidad": 1, "precio_unitario": 100.0},
        ]

    items_display_df = _build_items_dataframe(pd.DataFrame(st.session_state[items_state_key]))
    items_raw = st.data_editor(
        items_display_df,
        num_rows="dynamic",
        use_container_width=True,
        key="cotizacion_privada_items",
        column_config={
            "producto_servicio": st.column_config.TextColumn("Producto / Servicio", width="large", required=True),
            "cantidad": st.column_config.NumberColumn("Cantidad", min_value=0.0, step=1.0, required=True),
            "precio_unitario": st.column_config.NumberColumn(
                "Precio unitario", min_value=0.0, step=10.0, format="$%0.2f", required=True
            ),
            "importe": st.column_config.NumberColumn(
                "Subtotal", format="$%0.2f", disabled=True
            ),
        },
        disabled=["importe"],
        hide_index=True,
    )

    items_df = _build_items_dataframe(pd.DataFrame(items_raw))
    st.session_state[items_state_key] = items_df[
        ["producto_servicio", "cantidad", "precio_unitario"]
    ].to_dict(orient="records")
    subtotal = float(items_df["importe"].sum())
    impuesto_valor = subtotal * (float(impuesto_pct) / 100.0)
    total = subtotal + impuesto_valor

    st.markdown(
        f"**Resumen:** Subtotal {_format_money(subtotal)} | Impuesto ({impuesto_pct:.2f}%) {_format_money(impuesto_valor)} | Total {_format_money(total)}"
    )
    detalles_extra = st.text_area(
        "Detalles adicionales",
        height=90,
        key="cot_detalles_extra",
        placeholder="Agrega notas adicionales para la cotizacion.",
    )


    st.markdown("### Presupuesto interno")
    if presupuesto_state_key not in st.session_state:
        st.session_state[presupuesto_state_key] = [
            {"producto_servicio": "Detalle", "cantidad": 1, "precio_unitario": 0.0},
        ]

    presupuesto_display_df = _build_items_dataframe(pd.DataFrame(st.session_state[presupuesto_state_key]))
    presupuesto_raw = st.data_editor(
        presupuesto_display_df,
        num_rows="dynamic",
        use_container_width=True,
        key="cotizacion_presupuesto_items",
        column_config={
            "producto_servicio": st.column_config.TextColumn("Producto / Servicio", width="large", required=True),
            "cantidad": st.column_config.NumberColumn("Cantidad", min_value=0.0, step=1.0, required=True),
            "precio_unitario": st.column_config.NumberColumn(
                "Precio unitario", min_value=0.0, step=10.0, format="$%0.2f", required=True
            ),
            "importe": st.column_config.NumberColumn(
                "Subtotal", format="$%0.2f", disabled=True
            ),
        },
        disabled=["importe"],
        hide_index=True,
    )

    presupuesto_df = _build_items_dataframe(pd.DataFrame(presupuesto_raw))
    st.session_state[presupuesto_state_key] = presupuesto_df[
        ["producto_servicio", "cantidad", "precio_unitario"]
    ].to_dict(orient="records")

    costo_interno = float(presupuesto_df["importe"].sum())
    col_p1, col_p2, col_p3, col_p4 = st.columns([1, 1, 1, 1])
    with col_p1:
        factor_ganancia = st.number_input(
            "Factor de ganancia",
            min_value=0.0,
            step=0.05,
            key="cot_presupuesto_factor",
        )
    with col_p2:
        tiempo_inversion = st.number_input(
            "Tiempo inversion→inicio de ejecucion (dias)",
            min_value=0.0,
            step=1.0,
            key="cot_presupuesto_t_inversion",
        )
        inversion_etapa_1 = st.number_input(
            "Inversion requerida (etapa inversion→inicio de ejecucion)",
            min_value=0.0,
            step=100.0,
            key="cot_presupuesto_inv_etapa_1",
        )
    with col_p3:
        tiempo_intermedio = st.number_input(
            "Tiempo inicio de ejecucion→presentacion (dias)",
            min_value=0.0,
            step=1.0,
            key="cot_presupuesto_t_intermedio",
        )
        inversion_etapa_intermedia = st.number_input(
            "Inversion requerida (etapa inicio de ejecucion→presentacion)",
            min_value=0.0,
            step=100.0,
            key="cot_presupuesto_inv_etapa_intermedia",
        )
    with col_p4:
        tiempo_cobro = st.number_input(
            "Tiempo presentacion→cobro (dias)",
            min_value=0.0,
            step=1.0,
            key="cot_presupuesto_t_cobro",
        )
        inversion_etapa_2 = st.number_input(
            "Inversion requerida (etapa presentacion→cobro)",
            min_value=0.0,
            step=100.0,
            key="cot_presupuesto_inv_etapa_2",
        )

    col_f1, col_f2 = st.columns([1, 1])
    with col_f1:
        financiamiento_tipo = st.selectbox(
            "Financiamiento",
            ["Dinero propio", "Prestamo"],
            key="cot_presupuesto_fin_tipo",
        )
    with col_f2:
        financiamiento_interes_pct = st.number_input(
            "Interes mensual (%)",
            min_value=0.0,
            step=0.1,
            key="cot_presupuesto_fin_interes",
        )

    tiempo_recuperacion = tiempo_inversion + tiempo_intermedio + tiempo_cobro
    tiempo_recuperacion_meses = tiempo_recuperacion / 30 if tiempo_recuperacion else 0.0
    costo_financiamiento = 0.0
    if financiamiento_tipo == "Prestamo":
        tasa_mensual = financiamiento_interes_pct / 100.0
        meses_etapa_1 = tiempo_inversion / 30 if tiempo_inversion else 0.0
        meses_etapa_intermedia = tiempo_intermedio / 30 if tiempo_intermedio else 0.0
        meses_etapa_2 = tiempo_cobro / 30 if tiempo_cobro else 0.0
        if inversion_etapa_1 <= 0 and inversion_etapa_intermedia <= 0 and inversion_etapa_2 <= 0:
            costo_financiamiento = costo_interno * tasa_mensual * tiempo_recuperacion_meses
        else:
            costo_financiamiento = (
                (inversion_etapa_1 * tasa_mensual * meses_etapa_1)
                + (inversion_etapa_intermedia * tasa_mensual * meses_etapa_intermedia)
                + (inversion_etapa_2 * tasa_mensual * meses_etapa_2)
            )
    costo_total_final = costo_interno + costo_financiamiento
    precio_cotizar = costo_total_final * factor_ganancia
    ganancia = precio_cotizar - costo_total_final
    ganancia_neta = precio_cotizar - costo_total_final

    st.markdown(
        f"**Resumen presupuesto:** Costo interno {_format_money(costo_interno)} | "
        f"Base final (interno + financiamiento) {_format_money(costo_total_final)} | "
        f"Precio a cotizar {_format_money(precio_cotizar)} | "
        f"Ganancia {_format_money(ganancia)} | "
        f"Inversion etapa 1 {_format_money(inversion_etapa_1)} | "
        f"Inversion etapa 2 {_format_money(inversion_etapa_intermedia)} | "
        f"Inversion etapa 3 {_format_money(inversion_etapa_2)} | "
        f"Costo financiamiento {_format_money(costo_financiamiento)} | "
        f"Ganancia neta {_format_money(ganancia_neta)} | "
        f"Tiempo recuperacion {tiempo_recuperacion:.0f} dias (~{tiempo_recuperacion_meses:.1f} meses)"
    )

    condiciones = {
        "Vigencia": vigencia or "-",
        "Condicion de pago": forma_pago or "-",
        "Entrega": entrega or "-",
        "Lugar de entrega": lugar_entrega or "-",
    }
    st.markdown("### Vista previa (Excel final)")
    excel_preview_name = f"{numero_cot}.xlsx"
    excel_preview_bytes = None
    excel_preview_error = None
    try:
        excel_preview_bytes = _build_standard_quote_excel(
            empresa=empresa,
            numero_cot=numero_cot,
            fecha_cot=fecha_cot,
            cliente=cliente,
            direccion=direccion,
            cliente_ruc=cliente_ruc,
            cliente_dv=cliente_dv,
            items_df=items_df,
            impuesto_pct=impuesto_pct,
            condiciones=condiciones,
            detalles_extra=detalles_extra,
        )
        st.session_state["cot_std_excel_preview_bytes"] = excel_preview_bytes
        st.session_state["cot_std_excel_preview_name"] = excel_preview_name
    except Exception as exc:
        excel_preview_error = str(exc)
        st.warning(f"No se pudo construir la vista previa Excel: {exc}")

    if excel_preview_bytes:
        preview = _extract_standard_excel_preview(excel_preview_bytes)
        c1, c2, c3 = st.columns(3)
        with c1:
            st.caption("NÚMERO")
            st.write(preview.get("numero") or numero_cot)
            st.caption("CLIENTE")
            st.write(preview.get("cliente") or "-")
        with c2:
            st.caption("FECHA")
            st.write(preview.get("fecha") or "-")
            st.caption("RUC/DV")
            st.write(preview.get("ruc_dv") or "-")
        with c3:
            st.caption("TÍTULO")
            st.write(preview.get("titulo") or "-")

        preview_items_df = pd.DataFrame(preview.get("items") or [])
        if not preview_items_df.empty:
            st.dataframe(
                preview_items_df,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Costo Unitario": st.column_config.NumberColumn("Costo Unitario", format="$%0.2f"),
                    "Total": st.column_config.NumberColumn("Total", format="$%0.2f"),
                },
            )
        st.markdown(
            f"**Totales en Excel:** Subtotal {_format_money(preview.get('subtotal', 0.0))} | "
            f"Impuesto {_format_money(preview.get('impuesto', 0.0))} | "
            f"Total {_format_money(preview.get('total', 0.0))}"
        )
        if preview.get("condiciones"):
            st.caption("Condiciones colocadas en plantilla:")
            for line in preview["condiciones"][:8]:
                st.write(f"- {line}")

        st.download_button(
            "Descargar Excel de cotización",
            data=excel_preview_bytes,
            file_name=excel_preview_name,
            mime=_guess_mime_from_filename(excel_preview_name),
        )

    if st.button("Guardar cotización en Sheets/Drive"):
        if sheet_error or not sheet_id:
            st.error("No hay conexión a Google Sheets para guardar la cotización.")
        else:
            try:
                if client is None or creds is None:
                    client, creds = get_client()
                _ensure_cotizaciones_sheet(client, sheet_id)
                df_write = _normalize_cotizaciones_df(read_worksheet(client, sheet_id, SHEET_NAME_COT))

                now = datetime.now().isoformat(timespec="seconds")
                row_id = edit_row.get("id") if edit_row else uuid.uuid4().hex
                created_at = edit_row.get("created_at") if edit_row else now

                items_json = json.dumps(
                    items_df[["producto_servicio", "cantidad", "precio_unitario"]].to_dict(orient="records"),
                    ensure_ascii=False,
                )
                presupuesto_items_json = json.dumps(
                    presupuesto_df[["producto_servicio", "cantidad", "precio_unitario"]].to_dict(orient="records"),
                    ensure_ascii=False,
                )
                condiciones_json = json.dumps(condiciones, ensure_ascii=False)

                drive_file_id = edit_row.get("drive_file_id") if edit_row else ""
                drive_file_name = edit_row.get("drive_file_name") if edit_row else ""
                drive_file_url = edit_row.get("drive_file_url") if edit_row else ""
                drive_folder = edit_row.get("drive_folder") if edit_row else ""
                presupuesto_drive_file_id = edit_row.get("presupuesto_drive_file_id") if edit_row else ""
                presupuesto_drive_file_name = edit_row.get("presupuesto_drive_file_name") if edit_row else ""
                presupuesto_drive_file_url = edit_row.get("presupuesto_drive_file_url") if edit_row else ""
                tipo_cotizacion = (
                    edit_row.get("tipo_cotizacion")
                    if edit_row and edit_row.get("tipo_cotizacion")
                    else ("Estandar" if active_tab == "Cotizacion - Estandar" else "Panama Compra")
                )
                item_names_for_desc = [
                    str(val).strip()
                    for val in items_df.get("producto_servicio", pd.Series(dtype=str)).tolist()
                    if str(val).strip()
                ][:5]
                descripcion_corta = _generate_quote_short_description(
                    tipo_cotizacion=tipo_cotizacion,
                    empresa=empresa,
                    cliente=cliente,
                    detalles=detalles_extra,
                    items=item_names_for_desc,
                )
                excel_filename = _quote_excel_filename(numero_cot, descripcion_corta)
                if excel_preview_bytes:
                    excel_bytes = excel_preview_bytes
                else:
                    excel_bytes = _build_standard_quote_excel(
                        empresa=empresa,
                        numero_cot=numero_cot,
                        fecha_cot=fecha_cot,
                        cliente=cliente,
                        direccion=direccion,
                        cliente_ruc=cliente_ruc,
                        cliente_dv=cliente_dv,
                        items_df=items_df,
                        impuesto_pct=impuesto_pct,
                        condiciones=condiciones,
                        detalles_extra=detalles_extra,
                    )
                st.session_state["cot_std_excel_preview_bytes"] = excel_bytes
                st.session_state["cot_std_excel_preview_name"] = excel_filename

                if creds is not None:
                    drive = _get_drive_client(creds)
                    _, folders = _get_drive_folders(drive)
                    folder_id = folders.get(empresa)
                    if folder_id:
                        upload = _upload_drive_binary(
                            drive,
                            folder_id,
                            excel_filename,
                            excel_bytes,
                            _guess_mime_from_filename(excel_filename),
                            existing_file_id=drive_file_id or None,
                        )
                        drive_file_id = upload.get("id", drive_file_id)
                        drive_file_name = upload.get("name", excel_filename)
                        drive_folder = folder_id
                        if drive_file_id:
                            drive_file_url = f"https://drive.google.com/file/d/{drive_file_id}/view"
                        presupuesto_html = _build_budget_html(
                            empresa=empresa,
                            numero=numero_cot,
                            fecha_cot=fecha_cot,
                            presupuesto_df=presupuesto_df,
                            costo_interno=costo_interno,
                            factor_ganancia=factor_ganancia,
                            precio_cotizar=precio_cotizar,
                            ganancia=ganancia,
                            financiamiento_tipo=financiamiento_tipo,
                            financiamiento_interes_pct=financiamiento_interes_pct,
                            costo_financiamiento=costo_financiamiento,
                            ganancia_neta=ganancia_neta,
                            tiempo_inversion=tiempo_inversion,
                            inversion_etapa_1=inversion_etapa_1,
                            tiempo_inicio_ejecucion_presentacion=tiempo_intermedio,
                            inversion_etapa_intermedia=inversion_etapa_intermedia,
                            tiempo_cobro=tiempo_cobro,
                            inversion_etapa_2=inversion_etapa_2,
                        )
                        presupuesto_filename = f"Presupuesto_{numero_cot}.html"
                        presupuesto_upload = _upload_quote_html(
                            drive,
                            folder_id,
                            presupuesto_filename,
                            presupuesto_html,
                            existing_file_id=presupuesto_drive_file_id or None,
                        )
                        presupuesto_drive_file_id = presupuesto_upload.get("id", presupuesto_drive_file_id)
                        presupuesto_drive_file_name = presupuesto_upload.get("name", presupuesto_filename)
                        if presupuesto_drive_file_id:
                            presupuesto_drive_file_url = (
                                f"https://drive.google.com/file/d/{presupuesto_drive_file_id}/view"
                            )

                row = {
                    "id": row_id,
                    "numero_cotizacion": numero_cot,
                    "prefijo": prefijo,
                    "secuencia": seq,
                    "empresa": empresa,
                    "tipo_cotizacion": tipo_cotizacion,
                    "descripcion_corta": descripcion_corta,
                    "cliente_nombre": cliente,
                    "cliente_direccion": direccion,
                    "cliente_ruc": cliente_ruc,
                    "cliente_dv": cliente_dv,
                    "fecha_cotizacion": fecha_cot.isoformat(),
                    "created_at": created_at,
                    "updated_at": now,
                    "moneda": "USD",
                    "subtotal": subtotal,
                    "impuesto_pct": impuesto_pct,
                    "impuesto_monto": impuesto_valor,
                    "total": total,
                    "items_json": items_json,
                    "items_resumen": _items_resumen(items_df),
                    "detalles_extra": detalles_extra,
                    "presupuesto_items_json": presupuesto_items_json,
                    "presupuesto_subtotal": costo_interno,
                    "presupuesto_factor_ganancia": factor_ganancia,
                    "presupuesto_precio_cotizar": precio_cotizar,
                    "presupuesto_ganancia": ganancia,
                    "presupuesto_financiamiento_tipo": financiamiento_tipo,
                    "presupuesto_financiamiento_interes_pct": financiamiento_interes_pct,
                    "presupuesto_costo_financiamiento": costo_financiamiento,
                    "presupuesto_ganancia_neta": ganancia_neta,
                    "presupuesto_t_inversion_presentacion": tiempo_inversion,
                    "presupuesto_inversion_etapa_1": inversion_etapa_1,
                    "presupuesto_t_inicio_ejecucion_presentacion": tiempo_intermedio,
                    "presupuesto_inversion_etapa_intermedia": inversion_etapa_intermedia,
                    "presupuesto_t_presentacion_cobro": tiempo_cobro,
                    "presupuesto_inversion_etapa_2": inversion_etapa_2,
                    "presupuesto_t_recuperacion": tiempo_recuperacion,
                    "condiciones_json": condiciones_json,
                    "vigencia": vigencia,
                    "forma_pago": forma_pago,
                    "entrega": entrega,
                    "lugar_entrega": lugar_entrega,
                    "estado": edit_row.get("estado", "vigente") if edit_row else "vigente",
                    "notas": edit_row.get("notas", "") if edit_row else "",
                    "drive_file_id": drive_file_id,
                    "drive_file_name": drive_file_name,
                    "drive_file_url": drive_file_url,
                    "drive_folder": drive_folder,
                    "presupuesto_drive_file_id": presupuesto_drive_file_id,
                    "presupuesto_drive_file_name": presupuesto_drive_file_name,
                    "presupuesto_drive_file_url": presupuesto_drive_file_url,
                }

                if edit_row and row_id in df_write["id"].values:
                    idx = df_write.index[df_write["id"] == row_id][0]
                    for col in COT_COLUMNS:
                        df_write.at[idx, col] = row.get(col, "")
                else:
                    df_write = pd.concat([df_write, pd.DataFrame([row])], ignore_index=True)

                write_worksheet(client, sheet_id, SHEET_NAME_COT, df_write)
                st.session_state["cotizaciones_cache_token"] = uuid.uuid4().hex
                _clear_edit_state()
                st.success("Cotización guardada correctamente.")
            except Exception as exc:
                st.error(f"No se pudo guardar la cotización: {exc}")

if active_tab == "Historial de cotizaciones":
    if sheet_error:
        st.warning(sheet_error)
    else:
        if cot_df.empty:
            st.info("Aún no hay cotizaciones registradas.")
        else:
            desc_col = "descripcion_corta"
            missing_mask = cot_df[desc_col].fillna("").astype(str).str.strip().eq("")
            missing_count = int(missing_mask.sum())
            if missing_count:
                st.caption(f"Cotizaciones sin descripción corta: {missing_count}")
            auto_desc_key = "cot_hist_desc_autofill_once_done"
            if missing_count and not st.session_state.get(auto_desc_key, False):
                try:
                    if client is None:
                        client, creds = get_client()
                    df_write = cot_df.copy()
                    pending_idx = df_write.index[
                        df_write[desc_col].fillna("").astype(str).str.strip().eq("")
                    ].tolist()
                    total_pending = len(pending_idx)
                    progress = st.progress(0.0, text="Generando descripciones automáticamente...")
                    for pos, idx in enumerate(pending_idx, start=1):
                        row_data = df_write.loc[idx].to_dict()
                        df_write.at[idx, desc_col] = _generate_description_for_row(row_data)
                        progress.progress(
                            pos / total_pending,
                            text=f"Generando descripciones automáticamente... {pos}/{total_pending}",
                        )
                    write_worksheet(client, sheet_id, SHEET_NAME_COT, _normalize_cotizaciones_df(df_write))
                    st.session_state["cotizaciones_cache_token"] = uuid.uuid4().hex
                    st.session_state[auto_desc_key] = True
                    st.success(f"Descripciones generadas automáticamente: {total_pending}")
                    st.rerun()
                except Exception as exc:
                    st.session_state[auto_desc_key] = True
                    st.error(f"No se pudieron generar descripciones automáticamente: {exc}")

            search_text = st.text_input(
                "Buscar cotizaciones",
                key="cot_hist_search",
                placeholder="Número, cliente, descripción, empresa, tipo...",
            ).strip()
            filtered_df = cot_df.copy()
            if search_text:
                search_cols = [
                    "numero_cotizacion",
                    "cliente_nombre",
                    "descripcion_corta",
                    "empresa",
                    "tipo_cotizacion",
                    "detalles_extra",
                    "items_resumen",
                ]
                haystack = filtered_df[search_cols].fillna("").astype(str).agg(" ".join, axis=1)
                filtered_df = filtered_df[haystack.str.contains(search_text, case=False, na=False)]

            display_cols = [
                "numero_cotizacion",
                "empresa",
                "fecha_cotizacion",
                "cliente_nombre",
                "descripcion_corta",
                "total",
                "estado",
            ]
            if filtered_df.empty:
                st.info("No hay cotizaciones que coincidan con la búsqueda.")
                st.stop()
            st.dataframe(filtered_df[display_cols], use_container_width=True)

            opciones = filtered_df["id"].tolist()
            if (
                "cot_hist_selected_id" not in st.session_state
                or st.session_state["cot_hist_selected_id"] not in opciones
            ):
                st.session_state["cot_hist_selected_id"] = opciones[0]

            def _label(opt):
                row = filtered_df[filtered_df["id"] == opt].iloc[0]
                descripcion = str(row.get("descripcion_corta") or "").strip()
                base = f"{row.get('numero_cotizacion', '')} · {row.get('cliente_nombre', '')}"
                return f"{base} · {descripcion}" if descripcion else base

            selected_id = st.selectbox(
                "Selecciona una cotización",
                opciones,
                format_func=_label,
                key="cot_hist_selected_id",
            )
            sel_row = cot_df[cot_df["id"] == selected_id].iloc[0].to_dict()

            st.markdown("#### Detalle")
            st.write(
                {
                    "Número": sel_row.get("numero_cotizacion"),
                    "Empresa": sel_row.get("empresa"),
                    "Cliente": sel_row.get("cliente_nombre"),
                    "Descripción": sel_row.get("descripcion_corta"),
                    "RUC": sel_row.get("cliente_ruc"),
                    "DV": sel_row.get("cliente_dv"),
                    "Fecha": sel_row.get("fecha_cotizacion"),
                    "Total": sel_row.get("total"),
                }
            )
            detalles_hist = str(sel_row.get("detalles_extra") or "").strip()
            if detalles_hist:
                st.text_area(
                    "Detalles adicionales",
                    value=detalles_hist,
                    height=90,
                    disabled=True,
                )

            presupuesto_items = []
            try:
                presupuesto_items = json.loads(sel_row.get("presupuesto_items_json") or "[]")
                if not isinstance(presupuesto_items, list):
                    presupuesto_items = []
            except Exception:
                presupuesto_items = []
            if presupuesto_items:
                st.markdown("#### Presupuesto interno")
                pres_df = _build_items_dataframe(pd.DataFrame(presupuesto_items))
                st.dataframe(
                    pres_df[["producto_servicio", "cantidad", "precio_unitario", "importe"]],
                    use_container_width=True,
                    hide_index=True,
                )
                try:
                    pres_subtotal = float(sel_row.get("presupuesto_subtotal") or 0)
                except (TypeError, ValueError):
                    pres_subtotal = 0.0
                try:
                    pres_factor = float(sel_row.get("presupuesto_factor_ganancia") or 0)
                except (TypeError, ValueError):
                    pres_factor = 0.0
                try:
                    pres_precio = float(sel_row.get("presupuesto_precio_cotizar") or 0)
                except (TypeError, ValueError):
                    pres_precio = 0.0
                try:
                    pres_ganancia = float(sel_row.get("presupuesto_ganancia") or 0)
                except (TypeError, ValueError):
                    pres_ganancia = 0.0
                pres_fin_tipo = sel_row.get("presupuesto_financiamiento_tipo") or "Dinero propio"
                try:
                    pres_fin_interes = float(sel_row.get("presupuesto_financiamiento_interes_pct") or 0)
                except (TypeError, ValueError):
                    pres_fin_interes = 0.0
                try:
                    pres_costo_fin = float(sel_row.get("presupuesto_costo_financiamiento") or 0)
                except (TypeError, ValueError):
                    pres_costo_fin = 0.0
                try:
                    pres_ganancia_neta = float(sel_row.get("presupuesto_ganancia_neta") or 0)
                except (TypeError, ValueError):
                    pres_ganancia_neta = 0.0
                try:
                    pres_inv_etapa_1 = float(sel_row.get("presupuesto_inversion_etapa_1") or 0)
                except (TypeError, ValueError):
                    pres_inv_etapa_1 = 0.0
                try:
                    pres_inv_etapa_intermedia = float(sel_row.get("presupuesto_inversion_etapa_intermedia") or 0)
                except (TypeError, ValueError):
                    pres_inv_etapa_intermedia = 0.0
                try:
                    pres_inv_etapa_2 = float(sel_row.get("presupuesto_inversion_etapa_2") or 0)
                except (TypeError, ValueError):
                    pres_inv_etapa_2 = 0.0
                try:
                    pres_t_rec = float(sel_row.get("presupuesto_t_recuperacion") or 0)
                except (TypeError, ValueError):
                    pres_t_rec = 0.0
                pres_t_rec_meses = pres_t_rec / 30 if pres_t_rec else 0.0
                st.markdown(
                    f"**Resumen presupuesto:** Costo interno {_format_money(pres_subtotal)} | "
                    f"Factor {pres_factor:.2f} | Precio a cotizar {_format_money(pres_precio)} | "
                    f"Ganancia {_format_money(pres_ganancia)} | "
                    f"Financiamiento {pres_fin_tipo} ({pres_fin_interes:.2f}% mensual) | "
                    f"Inversion etapa 1 {_format_money(pres_inv_etapa_1)} | "
                    f"Inversion etapa 2 {_format_money(pres_inv_etapa_intermedia)} | "
                    f"Inversion etapa 3 {_format_money(pres_inv_etapa_2)} | "
                    f"Costo financiamiento {_format_money(pres_costo_fin)} | "
                    f"Ganancia neta {_format_money(pres_ganancia_neta)} | "
                    f"Tiempo recuperacion {pres_t_rec:.0f} dias (~{pres_t_rec_meses:.1f} meses)"
                )

            col_a, col_b, col_c, col_d = st.columns(4)
            with col_a:
                if st.button("Cargar en formulario"):
                    st.session_state[PENDING_EDIT_KEY] = selected_id
                    tipo = sel_row.get("tipo_cotizacion")
                    target_tab = (
                        "Cotizaci?n - Panam? Compra" if tipo == "Panama Compra" else "Cotizacion - Estandar"
                    )
                    st.session_state[PENDING_TAB_KEY] = target_tab
                    st.success("Cotizaci?n cargada en el formulario de edici?n.")
                    st.rerun()
            with col_b:
                if st.button("Duplicar"):
                    st.session_state[PENDING_DUPLICATE_KEY] = selected_id
                    tipo = sel_row.get("tipo_cotizacion")
                    target_tab = (
                        "Cotizaci?n - Panam? Compra" if tipo == "Panama Compra" else "Cotizacion - Estandar"
                    )
                    st.session_state[PENDING_TAB_KEY] = target_tab
                    st.success("Cotizaci?n duplicada en el formulario.")
                    st.rerun()
            with col_c:
                delete_key = f"delete_{selected_id}"
                if st.button("Eliminar"):
                    st.session_state[delete_key] = True
                if st.session_state.get(delete_key):
                    if st.button("Confirmar eliminaci?n"):
                        try:
                            if client is None:
                                client, creds = get_client()
                            df_write = cot_df[cot_df["id"] != selected_id].copy()
                            write_worksheet(client, sheet_id, SHEET_NAME_COT, df_write)
                            if sel_row.get("drive_file_id") and creds is not None:
                                drive = _get_drive_client(creds)
                                drive.files().delete(
                                    fileId=sel_row["drive_file_id"],
                                    supportsAllDrives=True,
                                ).execute()
                            st.session_state["cotizaciones_cache_token"] = uuid.uuid4().hex
                            st.success("Cotizaci?n eliminada.")
                            st.rerun()
                        except Exception as exc:
                            st.error(f"No se pudo eliminar: {exc}")
            with col_d:
                download_key = f"download_{selected_id}"
                if sel_row.get("drive_file_id"):
                    if st.button("Preparar descarga"):
                        try:
                            if creds is None:
                                client, creds = get_client()
                            drive = _get_drive_client(creds)
                            file_bytes = _download_drive_file(drive, sel_row["drive_file_id"])
                            st.session_state[download_key] = file_bytes
                        except Exception as exc:
                            st.error(f"No se pudo descargar: {exc}")
                    if st.session_state.get(download_key):
                        out_name = sel_row.get("drive_file_name") or f"{sel_row.get('numero_cotizacion')}.xlsx"
                        st.download_button(
                            "Descargar archivo",
                            data=st.session_state[download_key],
                            file_name=out_name,
                            mime=_guess_mime_from_filename(out_name),
                        )
                if sel_row.get("drive_file_url"):
                    st.link_button("Abrir en Drive", sel_row["drive_file_url"])
                if sel_row.get("presupuesto_drive_file_url"):
                    st.link_button("Abrir presupuesto", sel_row["presupuesto_drive_file_url"])
