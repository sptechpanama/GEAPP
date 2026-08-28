from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Iterable, Mapping, Optional

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image as PILImage


SP_COMPANY_NAME = "SP Tech Solutions S.A."


@dataclass(frozen=True)
class LPCompanyProfile:
    name: str
    code: str
    legal_name: str
    representative: str
    representative_id: str
    ruc: str
    dv: str
    operations_notice: str
    address: str
    phone: str
    header_filename: str
    signature_filename: str
    header_crop_bottom: float
    header_mode: str = "image"
    logo_filename: str = ""
    header_lines: tuple[str, ...] = ()


LP_COMPANY_PROFILES: dict[str, LPCompanyProfile] = {
    "RS Engineering": LPCompanyProfile(
        name="RS Engineering",
        code="RS",
        legal_name="RS ENGINEERING",
        representative="Rodrigo Jesús Sánchez Prado",
        representative_id="9-740-624",
        ruc="9-740-624",
        dv="80",
        operations_notice="9-740-624-2017-549317",
        address=(
            "Provincia de Veraguas, distrito de Santiago, corregimiento de Santiago "
            "(Cabecera), calle 1, casa 71, Urbanización Altos de Miramar"
        ),
        phone="6847-5616",
        header_filename="encabezado.png",
        signature_filename="firma.png",
        header_crop_bottom=0.22,
    ),
    "RIR Medical": LPCompanyProfile(
        name="RIR Medical",
        code="RIR",
        legal_name="RIR MEDICAL ENGINEERING",
        representative="Rodrigo Jesús Sánchez Prado",
        representative_id="9-740-624",
        ruc="155750585-2-2024",
        dv="40",
        operations_notice="55750585-2-2024-2024-574365876",
        address=(
            "Provincia de Veraguas, distrito de Santiago, corregimiento de Santiago "
            "(Cabecera), barriada Urracá, calle 11A, edificio CYPSA, planta alta, local 2"
        ),
        phone="6847-5616",
        header_filename="encabezado_rir.png",
        signature_filename="firma.png",
        header_crop_bottom=0.25,
        header_mode="dynamic",
        logo_filename="Logo RIR Medical.png",
        header_lines=(
            "RUC: 155750585-2-2024 | DV: 40",
            "Santiago, Veraguas · Urracá",
            "Tel. +507 6847-5616",
        ),
    ),
    SP_COMPANY_NAME: LPCompanyProfile(
        name=SP_COMPANY_NAME,
        code="SP",
        legal_name="SP TECH SOLUTIONS, S.A.",
        representative="Irvin Jesus Sanchez Prado",
        representative_id="9-749-1629",
        ruc="155767402-2-2025",
        dv="25",
        operations_notice="155767402-2-2025-2025-574431318",
        address=(
            "PH Bonanza Plaza, piso 4, oficina 4B, calle 41 Este, corregimiento de "
            "Bella Vista, distrito y provincia de Panamá"
        ),
        phone="262-4515",
        header_filename="encabezado_sp.png",
        signature_filename="firma_irvin_sp.png",
        header_crop_bottom=0.25,
    ),
}


CLAUSE_HEADING_RE = re.compile(
    r"^\s*(PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|S[ÉE]PTIMA|OCTAVA|NOVENA|D[ÉE]CIMA)\b",
    flags=re.IGNORECASE,
)


def get_lp_company_profile(company_name: str) -> LPCompanyProfile:
    normalized = str(company_name or "").strip()
    if normalized == "SP Engineering":
        normalized = SP_COMPANY_NAME
    return LP_COMPANY_PROFILES.get(normalized, LP_COMPANY_PROFILES["RS Engineering"])


def _iter_inner_paragraphs(doc: Document) -> Iterable[Paragraph]:
    """Incluye tablas y cuadros de texto sin desmontar dibujos contenedores."""

    roots = [doc.element.body]
    for section in doc.sections:
        roots.extend(
            [
                section.header._element,
                section.first_page_header._element,
                section.even_page_header._element,
                section.footer._element,
                section.first_page_footer._element,
                section.even_page_footer._element,
            ]
        )

    for root in roots:
        for element in root.iter(qn("w:p")):
            # Un párrafo exterior puede contener un cuadro de texto con otros
            # párrafos. Editarlo eliminaría el dibujo; se procesan solo los hijos.
            if next(element.iterdescendants(tag=qn("w:p")), None) is not None:
                continue
            yield Paragraph(element, doc)


def _paragraph_is_in_table(paragraph: Paragraph) -> bool:
    parent = paragraph._p.getparent()
    while parent is not None:
        if parent.tag == qn("w:tc"):
            return True
        parent = parent.getparent()
    return False


def _paragraph_is_in_body(paragraph: Paragraph) -> bool:
    parent = paragraph._p.getparent()
    while parent is not None:
        if parent.tag == qn("w:body"):
            return True
        parent = parent.getparent()
    return False


def _replace_tokens(text: str, replacements: Mapping[str, str]) -> str:
    result = str(text or "")
    for key in sorted(replacements, key=len, reverse=True):
        result = result.replace(str(key), str(replacements[key] or ""))
    return result


def _replace_contractor_segment(text: str, contractor_identity: str) -> str:
    pattern = re.compile(
        r"(por una parte y por la otra,)\s+.*?\s+(quien en adelante se denominará EL CONTRATISTA)",
        flags=re.IGNORECASE,
    )
    return pattern.sub(rf"\1 {contractor_identity}, \2", text, count=1)


def _sp_identity_sentence(profile: LPCompanyProfile) -> str:
    return (
        f"{profile.representative}, panameño, mayor de edad, con cédula de identidad "
        f"personal {profile.representative_id}, actuando en mi condición de Representante "
        f"Legal de la empresa {profile.legal_name}, sociedad anónima debidamente constituida "
        f"conforme a las leyes de la República de Panamá, inscrita bajo RUC {profile.ruc}, "
        f"DV {profile.dv}, con domicilio en {profile.address}, teléfono {profile.phone}"
    )


def _sp_contractor_identity(profile: LPCompanyProfile) -> str:
    return (
        f"{profile.representative}, con cédula de identidad personal o pasaporte No. "
        f"{profile.representative_id}, actuando en nombre y representación de la empresa "
        f"{profile.legal_name}, sociedad anónima debidamente constituida con número de Aviso "
        f"de Operación No. {profile.operations_notice}, con domicilio en {profile.address}"
    )


def _rir_identity_sentence(profile: LPCompanyProfile) -> str:
    return (
        "Rodrigo Sánchez Prado varón panameño, mayor de edad, Ingeniero, soltero, con "
        f"cédula de identidad personal {profile.representative_id}, actuando en mi condición "
        f"de Representante Legal de la empresa {profile.legal_name}, sociedad debidamente "
        "constituida conforme a las leyes de la República de Panamá, debidamente inscrita "
        f"en el Registro Público de Panamá a Folio 155750585, con domicilio en {profile.address}, "
        f"teléfono {profile.phone}"
    )


def _apply_company_identity(text: str, profile: LPCompanyProfile) -> str:
    source = str(text or "")
    if not source:
        return source

    if profile.code == "SP":
        if source.startswith("En cumplimiento de lo establecido") and "el suscrito" in source:
            prefix, _, tail = source.partition("el suscrito")
            marker = "declaro BAJO LA GRAVEDAD DE JURAMENTO lo siguiente:"
            if marker in tail:
                source = f"{prefix}el suscrito {_sp_identity_sentence(profile)}, {marker}"
            elif "declaro lo siguiente:" in tail:
                source = f"{prefix}el suscrito {_sp_identity_sentence(profile)}; declaro lo siguiente:"
        elif source.startswith("Entre los suscritos a saber;"):
            source = _replace_contractor_segment(source, _sp_contractor_identity(profile))

        source = _replace_tokens(
            source,
            {
                "Rodrigo Jesús Sánchez Prado": profile.representative,
                "Rodrigo Jesus Sanchez Prado": profile.representative,
                "Rodrigo Sánchez Prado": profile.representative,
                "Rodrigo Sanchez Prado": profile.representative,
                "9-740-624-2017-549317": profile.operations_notice,
                "RS ENGINEERING": profile.legal_name,
                "RS Engineering": profile.legal_name,
                "DV:80": f"DV:{profile.dv}",
                "9-740-624": profile.representative_id,
                "ruc: 9-740-624": f"RUC: {profile.ruc}",
            },
        )

        # La sustitución del RUC debe ocurrir después de la cédula para evitar
        # que el RUC de la plantilla base quede convertido en la nueva cédula.
        source = source.replace(
            f"RUC: {profile.representative_id}, DV:{profile.dv}",
            f"RUC: {profile.ruc}, DV:{profile.dv}",
        )
        source = source.replace(
            f"ruc: {profile.representative_id}, DV:{profile.dv}",
            f"RUC: {profile.ruc}, DV:{profile.dv}",
        )
        return source

    if profile.code == "RIR":
        # La única plantilla heredada de RS es No Incapacidad; se corrige su
        # identidad sin tocar el cuerpo jurídico de la declaración.
        if (
            source.startswith("En cumplimiento de lo establecido")
            and "artículo 24" in source
            and "RS ENGINEERING" in source
        ):
            prefix, _, tail = source.partition("el suscrito")
            marker = "declaro BAJO LA GRAVEDAD DE JURAMENTO lo siguiente:"
            if marker in tail:
                source = f"{prefix}el suscrito {_rir_identity_sentence(profile)}, {marker}"
        source = source.replace("RS ENGINEERING", profile.legal_name)
        return source

    return source


def _highlight_values(
    replacements: Mapping[str, str], profile: LPCompanyProfile
) -> list[str]:
    values = [
        str(replacements.get(key, "") or "").strip()
        for key in (
            "[Representante_legal_de_la_Entidad_Licitante]",
            "[cedula]",
            "[numero_de_acto]",
            "[entidad]",
            "[titulo]",
            "[lugar]",
            "[entrega]",
        )
    ]
    values.extend(
        [
            profile.representative,
            profile.representative_id,
            profile.legal_name,
            profile.ruc,
            profile.operations_notice,
        ]
    )
    unique = {value.casefold(): value for value in values if value}
    return sorted(unique.values(), key=len, reverse=True)


def _find_spans(text: str, values: Iterable[str]) -> list[tuple[int, int, bool]]:
    spans: list[tuple[int, int, bool]] = []
    lowered = text.casefold()
    for value in values:
        needle = str(value or "").strip()
        if not needle:
            continue
        start = 0
        folded = needle.casefold()
        while True:
            index = lowered.find(folded, start)
            if index < 0:
                break
            spans.append((index, index + len(needle), False))
            start = index + len(needle)

    clause = CLAUSE_HEADING_RE.match(text)
    if clause:
        spans.append((clause.start(1), clause.end(1), True))
    return spans


def _rebuild_with_emphasis(paragraph: Paragraph, values: Iterable[str]) -> None:
    text = str(paragraph.text or "")
    spans = _find_spans(text, values)
    if not spans:
        return

    points = {0, len(text)}
    for start, end, _ in spans:
        points.update((start, end))
    cuts = sorted(points)
    paragraph.text = ""
    for index in range(len(cuts) - 1):
        start, end = cuts[index], cuts[index + 1]
        if start >= end:
            continue
        run = paragraph.add_run(text[start:end])
        if any(start < span_end and end > span_start for span_start, span_end, _ in spans):
            run.bold = True
        if any(
            underline and start < span_end and end > span_start
            for span_start, span_end, underline in spans
        ):
            run.underline = True


def _set_run_font(run, *, size: float, bold: Optional[bool] = None) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def _style_document(
    doc: Document,
    *,
    document_name: str,
    highlight_values: list[str],
) -> None:
    compact = "pacto_de_integridad" in document_name
    body_size = 9.15 if compact else 10.25
    table_size = 9.0 if compact else 9.5

    for section in doc.sections:
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(14)
        section.top_margin = Inches(1.42)
        section.bottom_margin = Inches(0.52)
        section.left_margin = Inches(0.62)
        section.right_margin = Inches(0.62)
        section.header_distance = Inches(0.12)
        section.footer_distance = Inches(0.25)

    paragraphs = list(_iter_inner_paragraphs(doc))
    first_body: Optional[Paragraph] = None
    for paragraph in paragraphs:
        if not _paragraph_is_in_body(paragraph):
            continue
        if paragraph.text.strip():
            first_body = paragraph
            break

    for paragraph in paragraphs:
        text = str(paragraph.text or "").strip()
        in_table = _paragraph_is_in_table(paragraph)
        _rebuild_with_emphasis(paragraph, highlight_values)

        pf = paragraph.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(1 if compact else (1.5 if in_table else 2.5))
        pf.line_spacing = 1.0 if compact else 1.05

        if paragraph is first_body:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_after = Pt(7)
            pf.keep_with_next = True
            for run in paragraph.runs:
                _set_run_font(run, size=13.0, bold=True)
            continue

        if text and len(text) >= 85 and not in_table:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif in_table:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        font_size = table_size if in_table else body_size
        for run in paragraph.runs:
            _set_run_font(run, size=font_size)


def _clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _set_table_bottom_border(table, *, color: str = "003B82") -> None:
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge_name in ("top", "left", "insideH", "insideV", "right"):
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "nil")
        borders.append(edge)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _add_dynamic_company_header(
    doc: Document,
    profile: LPCompanyProfile,
    assets_dir: Path,
) -> None:
    logo_path = assets_dir.parent / profile.logo_filename
    if not logo_path.exists():
        return

    for section in doc.sections:
        section.header.is_linked_to_previous = False
        header = section.header
        for old_table in list(header.tables):
            old_table._element.getparent().remove(old_table._element)

        spacer = header.paragraphs[0]
        _clear_paragraph(spacer)
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.line_spacing = 0.5

        table = header.add_table(rows=1, cols=2, width=Inches(7.25))
        table.autofit = False
        table.columns[0].width = Inches(3.55)
        table.columns[1].width = Inches(3.70)
        left, right = table.rows[0].cells
        left.width = Inches(3.55)
        right.width = Inches(3.70)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        left_paragraph = left.paragraphs[0]
        left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_paragraph.paragraph_format.space_after = Pt(0)
        left_paragraph.add_run().add_picture(str(logo_path), width=Inches(2.35))

        right_paragraph = right.paragraphs[0]
        right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_paragraph.paragraph_format.space_before = Pt(0)
        right_paragraph.paragraph_format.space_after = Pt(0)
        right_paragraph.paragraph_format.line_spacing = 1.0
        for index, line in enumerate(profile.header_lines):
            run = right_paragraph.add_run(line)
            _set_run_font(run, size=8.25, bold=index == 0)
            if index < len(profile.header_lines) - 1:
                run.add_break()
        _set_table_bottom_border(table)


def _add_company_header(doc: Document, profile: LPCompanyProfile, assets_dir: Path) -> None:
    if profile.header_mode == "dynamic":
        _add_dynamic_company_header(doc, profile, assets_dir)
        return

    header_path = assets_dir / profile.header_filename
    if not header_path.exists():
        return

    with PILImage.open(header_path) as image:
        width_px, height_px = image.size
    visible_height_px = max(height_px * (1.0 - profile.header_crop_bottom), 1.0)
    display_width = 5.8
    display_height = display_width / (width_px / visible_height_px)

    for section in doc.sections:
        section.header.is_linked_to_previous = False
        paragraph = section.header.paragraphs[0]
        _clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        shape = paragraph.add_run().add_picture(
            str(header_path),
            width=Inches(display_width),
            height=Inches(display_height),
        )
        blip_fill = shape._inline.graphic.graphicData.pic.blipFill
        source_rectangle = OxmlElement("a:srcRect")
        source_rectangle.set("b", str(round(profile.header_crop_bottom * 100000)))
        blip_fill.insert(1, source_rectangle)


def _body_has_drawing(doc: Document) -> bool:
    for element in doc.element.body.iter():
        if element.tag in {qn("w:drawing"), qn("w:pict")}:
            return True
    return False


def _ensure_signature_when_missing(
    doc: Document,
    *,
    document_name: str,
    signature_path: Path,
) -> None:
    if "_sf" in document_name or _body_has_drawing(doc) or not signature_path.exists():
        return
    for paragraph in _iter_inner_paragraphs(doc):
        if "________" not in str(paragraph.text or ""):
            continue
        _clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(signature_path), width=Inches(1.65))
        return


def _signature_canvas(signature_path: Path, width: int, height: int) -> PILImage.Image:
    with PILImage.open(signature_path) as source:
        signature = source.convert("RGBA")
    canvas = PILImage.new("RGBA", (width, height), (255, 255, 255, 0))
    max_width = max(int(width * 0.9), 1)
    max_height = max(int(height * 0.82), 1)
    scale = min(max_width / signature.width, max_height / signature.height)
    resized = signature.resize(
        (max(int(signature.width * scale), 1), max(int(signature.height * scale), 1)),
        PILImage.Resampling.LANCZOS,
    )
    position = ((width - resized.width) // 2, (height - resized.height) // 2)
    canvas.alpha_composite(resized, dest=position)
    return canvas


def _encode_image(image: PILImage.Image, suffix: str) -> bytes:
    output = BytesIO()
    if suffix in {".tif", ".tiff"}:
        image.save(output, format="TIFF", compression="tiff_lzw")
    else:
        image.save(output, format="PNG", optimize=True)
    return output.getvalue()


def _optimize_signature_media(doc_bytes: bytes, signature_path: Path) -> bytes:
    if not signature_path.exists():
        return doc_bytes

    source = BytesIO(doc_bytes)
    target = BytesIO()
    with zipfile.ZipFile(source, "r") as incoming, zipfile.ZipFile(
        target, "w", compression=zipfile.ZIP_DEFLATED
    ) as outgoing:
        for info in incoming.infolist():
            data = incoming.read(info.filename)
            suffix = Path(info.filename).suffix.lower()
            should_replace = (
                info.filename.startswith("word/media/")
                and suffix in {".png", ".tif", ".tiff"}
                and len(data) >= 100_000
            )
            if should_replace:
                try:
                    with PILImage.open(BytesIO(data)) as original:
                        ratio = original.width / max(original.height, 1)
                    width = 1200
                    height = max(int(width / max(ratio, 0.2)), 300)
                    canvas = _signature_canvas(signature_path, width, height)
                    data = _encode_image(canvas, suffix)
                except Exception:
                    # Si un recurso no es una imagen legible, se conserva tal cual.
                    pass
            outgoing.writestr(info, data)
    return target.getvalue()


def render_lp_document(
    template_path: Path,
    replacements: Mapping[str, str],
    *,
    company_name: str,
    document_name: str,
    assets_dir: Path,
) -> bytes:
    profile = get_lp_company_profile(company_name)
    document = Document(str(template_path))

    for paragraph in _iter_inner_paragraphs(document):
        original = str(paragraph.text or "")
        transformed = _apply_company_identity(original, profile)
        transformed = _replace_tokens(transformed, replacements)
        if transformed != original:
            paragraph.text = transformed

    highlights = _highlight_values(replacements, profile)
    signature_path = assets_dir / profile.signature_filename
    _ensure_signature_when_missing(
        document,
        document_name=document_name,
        signature_path=signature_path,
    )
    _style_document(document, document_name=document_name, highlight_values=highlights)
    _add_company_header(document, profile, assets_dir)

    output = BytesIO()
    document.save(output)
    return _optimize_signature_media(output.getvalue(), signature_path)
